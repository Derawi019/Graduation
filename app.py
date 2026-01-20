from flask import (
    Flask, render_template, request, jsonify,
    redirect, url_for, flash, session, Response, send_file
)
from ai_core.assistant import handle_request
from flask_sqlalchemy import SQLAlchemy
from flask_login import (
    LoginManager, UserMixin,
    login_user, logout_user,
    login_required, current_user
)
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from flask_mail import Mail, Message
from deep_translator import GoogleTranslator
from langdetect import detect
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from flask_limiter.errors import RateLimitExceeded
from authlib.integrations.flask_client import OAuth
from dotenv import load_dotenv
import ffmpeg
from transformers import pipeline

from werkzeug.utils import secure_filename
from flask import request, jsonify
import subprocess
import shutil
import tempfile
import whisper
import os
import uuid
import time
import logging
from logging.handlers import RotatingFileHandler
from datetime import datetime
import json
import csv
import io
from gtts import gTTS
import re
import concurrent.futures
from transformers import AutoModelForCausalLM, AutoTokenizer
import torch

# Defer heavyweight model loading until first use.
whisper_model = None
dialogpt_tokenizer = None
dialogpt_model = None
summarizer = None

app = Flask(__name__)

OPENAI_AVAILABLE = False

def get_whisper_model():
    global whisper_model
    if whisper_model is None:
        model_name = os.getenv("WHISPER_MODEL", "base")
        whisper_model = whisper.load_model(model_name)
    return whisper_model


def get_dialogpt():
    global dialogpt_tokenizer, dialogpt_model
    if dialogpt_tokenizer is None or dialogpt_model is None:
        dialogpt_tokenizer = AutoTokenizer.from_pretrained("microsoft/DialoGPT-medium")
        # Prefer safetensors to avoid torch.load vulnerability checks
        dialogpt_model = AutoModelForCausalLM.from_pretrained(
            "microsoft/DialoGPT-medium",
            use_safetensors=True
        )
    return dialogpt_tokenizer, dialogpt_model


def get_summarizer():
    global summarizer
    if summarizer is None:
        # Abstractive summarization with a small, fast instruction-tuned model.
        summarizer = pipeline("text2text-generation", model="google/flan-t5-small")
    return summarizer


def quick_summary(text: str, max_sentences: int = 3, max_words: int = 60) -> str:
    """
    Fast extractive summary for better responsiveness.
    """
    cleaned = text.strip()
    if not cleaned:
        return text

    words = cleaned.split()
    if len(words) <= max_words:
        return cleaned

    sentences = re.split(r'(?<=[.!?])\s+', cleaned)
    sentences = [s for s in sentences if s]
    if sentences:
        summary = " ".join(sentences[:max_sentences]).strip()
    else:
        summary = " ".join(words[:max_words]).strip()

    summary_words = summary.split()
    if len(summary_words) > max_words:
        summary = " ".join(summary_words[:max_words]).strip()

    return summary + "..."


def generate_huggingface_response(user_input):
    # ترميز النص المدخل
    tokenizer, model = get_dialogpt()
    input_ids = tokenizer.encode(user_input + tokenizer.eos_token, return_tensors="pt")

    # توليد الرد من النموذج
    bot_output = model.generate(input_ids, max_length=1000, pad_token_id=tokenizer.eos_token_id)

    # فك ترميز الرد الناتج
    chatbot_response = tokenizer.decode(bot_output[:, input_ids.shape[-1]:][0], skip_special_tokens=True)
    
    return chatbot_response



def convert_to_wav(input_path):
    """
    Convert audio file to WAV (16kHz mono) using ffmpeg
    """
    if not shutil.which("ffmpeg"):
        raise RuntimeError("ffmpeg is not installed or not in PATH")

    wav_path = input_path.rsplit(".", 1)[0] + ".wav"

    command = [
        "ffmpeg",
        "-y",
        "-i", input_path,
        "-ac", "1",
        "-ar", "16000",
        wav_path
    ]

    subprocess.run(
        command,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        check=True
    )

    return wav_path

def transcribe_audio(audio_path):
    """
    Transcribe audio using Whisper AI (Windows-safe)
    """
    if not os.path.exists(audio_path):
        raise FileNotFoundError(f"Audio file not found: {audio_path}")

    wav_path = None
    try:
        # 🔁 تحويل webm → wav
        wav_path = convert_to_wav(audio_path)

        model = get_whisper_model()
        result = model.transcribe(
            wav_path,
            fp16=False
        )
        return result.get("text", "").strip()

    finally:
        # تنظيف ملف wav
        if wav_path and os.path.exists(wav_path):
            try:
                os.remove(wav_path)
            except Exception:
                pass

# Load environment variables from .env file
load_dotenv()

# Configure pydub to use local ffmpeg if available
# Check for local ffmpeg in bin directory (installed without brew)
# IMPORTANT: Patch pydub BEFORE importing AudioSegment, because AudioSegment
# imports mediainfo_json at module level
_script_dir = os.path.dirname(os.path.abspath(__file__))
_local_ffmpeg = os.path.join(_script_dir, 'bin', 'ffmpeg')
_local_ffprobe = os.path.join(_script_dir, 'bin', 'ffprobe')

# Define patched functions at module level so they're accessible
_patched_mediainfo_json = None

# Patch pydub.utils.mediainfo_json BEFORE importing AudioSegment
if os.path.exists(_local_ffprobe) and os.access(_local_ffprobe, os.X_OK):
    # Also add to PATH for subprocess calls
    bin_dir = os.path.join(_script_dir, 'bin')
    current_path = os.environ.get('PATH', '')
    if bin_dir not in current_path:
        os.environ['PATH'] = f"{bin_dir}:{current_path}"
    
    # Monkey-patch pydub's get_prober_name and mediainfo_json to use absolute path
    # This must be done BEFORE importing AudioSegment
    import pydub.utils
    _original_get_prober_name = pydub.utils.get_prober_name
    def patched_get_prober_name():
        """Patched version that returns absolute path to ffprobe"""
        return _local_ffprobe
    pydub.utils.get_prober_name = patched_get_prober_name
    
    # Also patch mediainfo_json to be safe
    _original_mediainfo_json = pydub.utils.mediainfo_json
    def patched_mediainfo_json(file_path, read_ahead_limit=-1):
        """Patched version that uses absolute path to ffprobe"""
        from subprocess import Popen, PIPE
        # Use the same logic as original but with absolute path
        prober = _local_ffprobe
        command_args = ["-v", "info", "-show_format", "-show_streams"]
        try:
            command_args.append(pydub.utils.fsdecode(file_path))
            stdin_parameter = None
            stdin_data = None
        except TypeError:
            command_args.extend(["-read_ahead_limit", str(read_ahead_limit), "cache:pipe:0"])
            stdin_parameter = PIPE
            file, close_file = pydub.utils._fd_or_path_or_tempfile(file_path, 'rb', tempfile=False)
            file.seek(0)
            stdin_data = file.read()
            if close_file:
                file.close()
        
        command = [prober, '-of', 'json'] + command_args
        res = Popen(command, stdin=stdin_parameter, stdout=PIPE, stderr=PIPE)
        output, stderr = res.communicate(input=stdin_data)
        output = output.decode("utf-8", 'ignore')
        stderr = stderr.decode("utf-8", 'ignore')
        try:
            return json.loads(output)
        except json.decoder.JSONDecodeError:
            # If JSON parsing fails, raise with error message
            raise Exception(f"ffprobe failed: {stderr}")
    _patched_mediainfo_json = patched_mediainfo_json
    pydub.utils.mediainfo_json = patched_mediainfo_json

# Now import AudioSegment (it will use the patched mediainfo_json)
from pydub import AudioSegment
# Also patch it in AudioSegment module namespace since it imports mediainfo_json
import pydub.audio_segment
if _patched_mediainfo_json is not None and hasattr(pydub.audio_segment, 'mediainfo_json'):
    # Force patch it in audio_segment module
    pydub.audio_segment.mediainfo_json = _patched_mediainfo_json

# Configure AudioSegment to use local ffmpeg/ffprobe
if os.path.exists(_local_ffmpeg) and os.access(_local_ffmpeg, os.X_OK):
    AudioSegment.converter = _local_ffmpeg
    AudioSegment.ffmpeg = _local_ffmpeg
    if os.path.exists(_local_ffprobe) and os.access(_local_ffprobe, os.X_OK):
        AudioSegment.ffprobe = _local_ffprobe
    else:
        AudioSegment.ffprobe = None
    # Note: Logging will be initialized later, so we'll log this after Flask app is created

# Initialize Sentry for error tracking (optional - only if SENTRY_DSN is set)
sentry_dsn = os.getenv('SENTRY_DSN')
if sentry_dsn:
    import sentry_sdk
    from sentry_sdk.integrations.flask import FlaskIntegration
    from sentry_sdk.integrations.sqlalchemy import SqlalchemyIntegration
    
    sentry_sdk.init(
        dsn=sentry_dsn,
        integrations=[
            FlaskIntegration(),
            SqlalchemyIntegration(),
        ],
        # Set traces_sample_rate to 1.0 to capture 100%
        # of the transactions for performance monitoring.
        # We recommend adjusting this value in production.
        traces_sample_rate=float(os.getenv('SENTRY_TRACES_SAMPLE_RATE', '0.1')),
        # Set profiles_sample_rate to 1.0 to profile 100%
        # of sampled transactions.
        # We recommend adjusting this value in production.
        profiles_sample_rate=float(os.getenv('SENTRY_PROFILES_SAMPLE_RATE', '0.1')),
        # Set environment
        environment=os.getenv('FLASK_ENV', 'production'),
        # Release version (optional)
        release=os.getenv('SENTRY_RELEASE'),
        # Before send callback to filter events
        before_send=lambda event, hint: event if os.getenv('SENTRY_ENABLED', 'true').lower() == 'true' else None,
    )

app = Flask(__name__)

# Flask configuration from environment variables
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'dev-key-change-in-production-12345')
app.config['MAX_CONTENT_LENGTH'] = int(os.getenv('MAX_CONTENT_LENGTH', 16 * 1024 * 1024))  # Default: 16MB
app.config['UPLOAD_FOLDER'] = os.getenv('UPLOAD_FOLDER', tempfile.gettempdir())

# ==============================
# Database configuration
# ==============================
db_user = os.getenv('DB_USER', 'postgres')
db_password = os.getenv('DB_PASSWORD', 'postgres')
db_host = os.getenv('DB_HOST', '127.0.0.1')
if db_host == 'localhost':
    db_host = '127.0.0.1'
db_port = os.getenv('DB_PORT', '5432')
db_name = os.getenv('DB_NAME', 'translation_app')

database_url = os.getenv('DATABASE_URL')

if not database_url:
    database_url = (
        f'postgresql://{db_user}:{db_password}'
        f'@{db_host}:{db_port}/{db_name}'
    )

app.config['SQLALCHEMY_DATABASE_URI'] = database_url
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)


# Initialize Flask-Mail
mail = Mail(app)
# Get email configuration from environment
mail_server = os.getenv('MAIL_SERVER', 'smtp.gmail.com')
# Validate MAIL_SERVER - if it looks like an email, default to Gmail SMTP
if '@' in mail_server and '.' in mail_server.split('@')[1]:
    app.logger.warning(f'MAIL_SERVER appears to be an email address ({mail_server}). Using smtp.gmail.com instead.')
    mail_server = 'smtp.gmail.com'

app.config['MAIL_SERVER'] = mail_server
app.config['MAIL_PORT'] = int(os.getenv('MAIL_PORT', '587'))
app.config['MAIL_USE_TLS'] = os.getenv('MAIL_USE_TLS', 'true').lower() == 'true'
app.config['MAIL_USE_SSL'] = os.getenv('MAIL_USE_SSL', 'false').lower() == 'true'
app.config['MAIL_USERNAME'] = os.getenv('MAIL_USERNAME', '')
app.config['MAIL_PASSWORD'] = os.getenv('MAIL_PASSWORD', '')
app.config['MAIL_DEFAULT_SENDER'] = os.getenv('MAIL_DEFAULT_SENDER', os.getenv('MAIL_USERNAME', 'noreply@translationapp.com'))
app.config['MAIL_SUPPRESS_SEND'] = os.getenv('MAIL_SUPPRESS_SEND', 'false').lower() == 'true'  # For testing

# Log email configuration (without password)
if app.config['MAIL_USERNAME']:
    app.logger.info(f'📧 Email configured: {app.config["MAIL_SERVER"]}:{app.config["MAIL_PORT"]} (User: {app.config["MAIL_USERNAME"]})')
else:
    app.logger.warning('⚠️  Email not configured - verification emails will not be sent')

# Initialize Flask-Login
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'
login_manager.login_message = 'Please log in to access this page.'
login_manager.login_message_category = 'info'
# Flask-Login stores next URL in session with key '_next' (with underscore)

# Initialize OAuth
oauth = OAuth(app)
google = oauth.register(
    name='google',
    client_id=os.getenv('GOOGLE_CLIENT_ID', ''),
    client_secret=os.getenv('GOOGLE_CLIENT_SECRET', ''),
    server_metadata_url='https://accounts.google.com/.well-known/openid-configuration',
    client_kwargs={
        'scope': 'openid email profile'
    }
)

# Initialize rate limiter
limiter = Limiter(
    app=app,
    key_func=get_remote_address,  # Rate limit by IP address
    default_limits=["200 per day", "50 per hour"],  # Global limits
    storage_uri="memory://"  # Use in-memory storage (use Redis for production)
)

# Configure logging
log_level = os.getenv('LOG_LEVEL', 'INFO').upper()
log_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'logs')

# Create logs directory if it doesn't exist
if not os.path.exists(log_dir):
    os.makedirs(log_dir)

# Configure logging format
log_format = logging.Formatter(
    '%(asctime)s [%(levelname)s] %(name)s: %(message)s [%(pathname)s:%(lineno)d]'
)

# Set up file handler with rotation (10MB per file, keep 10 backups)
file_handler = RotatingFileHandler(
    os.path.join(log_dir, 'translation_app.log'),
    maxBytes=10 * 1024 * 1024,  # 10MB
    backupCount=10,
    encoding='utf-8'
)
file_handler.setFormatter(log_format)
file_handler.setLevel(getattr(logging, log_level, logging.INFO))

# Set up console handler for development
console_handler = logging.StreamHandler()
console_handler.setFormatter(log_format)
console_handler.setLevel(getattr(logging, log_level, logging.INFO))

# Configure app logger
app.logger.setLevel(getattr(logging, log_level, logging.INFO))
app.logger.addHandler(file_handler)

# Only add console handler in development
if os.getenv('FLASK_ENV') == 'development' or app.debug:
    app.logger.addHandler(console_handler)

# Log startup
app.logger.info('=' * 60)
app.logger.info('Translation App Starting')
app.logger.info(f'Log Level: {log_level}')
app.logger.info(f'Environment: {os.getenv("FLASK_ENV", "production")}')
app.logger.info(f'Database: {db_name}@{db_host}:{db_port}')
if sentry_dsn:
    app.logger.info('✅ Sentry error tracking: ENABLED')
else:
    app.logger.info('⚠️  Sentry error tracking: DISABLED (set SENTRY_DSN to enable)')
# Log ffmpeg and ffprobe configuration
if os.path.exists(_local_ffmpeg) and os.access(_local_ffmpeg, os.X_OK):
    try:
        import subprocess
        ffmpeg_version = subprocess.check_output([_local_ffmpeg, '-version'], stderr=subprocess.STDOUT, timeout=2).decode().split('\n')[0]
        app.logger.info(f'✅ ffmpeg: Using local installation ({_local_ffmpeg})')
        app.logger.info(f'   {ffmpeg_version}')
    except Exception:
        app.logger.info(f'✅ ffmpeg: Using local installation ({_local_ffmpeg})')
else:
    app.logger.warning('⚠️  ffmpeg: Not found in bin/ directory. Audio conversion may not work properly.')

if os.path.exists(_local_ffprobe) and os.access(_local_ffprobe, os.X_OK):
    try:
        import subprocess
        ffprobe_version = subprocess.check_output([_local_ffprobe, '-version'], stderr=subprocess.STDOUT, timeout=2).decode().split('\n')[0]
        app.logger.info(f'✅ ffprobe: Using local installation ({_local_ffprobe})')
        app.logger.info(f'   {ffprobe_version}')
    except Exception:
        app.logger.info(f'✅ ffprobe: Using local installation ({_local_ffprobe})')
else:
    app.logger.warning('⚠️  ffprobe: Not found in bin/ directory. Some audio formats may not work properly.')

# Log OpenAI configuration
if OPENAI_AVAILABLE:
    openai_api_key = os.getenv('OPENAI_API_KEY')
    if openai_api_key:
        # Show only first 10 and last 4 characters for security
        masked_key = f"{openai_api_key[:10]}...{openai_api_key[-4:]}" if len(openai_api_key) > 14 else "***"
        app.logger.info(f'✅ OpenAI: Library available, API key configured ({masked_key})')
    else:
        app.logger.warning('⚠️  OpenAI: Library available but OPENAI_API_KEY not found in environment')
else:
    app.logger.warning('⚠️  OpenAI: Library not available (install with: pip install openai)')
app.logger.info('=' * 60)

# Database Models
class User(UserMixin, db.Model):
    """User model for authentication"""
    __tablename__ = 'users'
    
    id = db.Column(db.Integer, primary_key=True)
    email = db.Column(db.String(120), unique=True, nullable=False, index=True)
    username = db.Column(db.String(80), unique=True, nullable=True)
    password_hash = db.Column(db.String(255), nullable=True)  # Nullable for OAuth users
    name = db.Column(db.String(100), nullable=True)
    picture = db.Column(db.String(255), nullable=True)  # Profile picture URL
    provider = db.Column(db.String(50), nullable=False, default='local')  # 'local' or 'google'
    provider_id = db.Column(db.String(255), nullable=True, index=True)  # OAuth provider user ID
    created_at = db.Column(db.DateTime, default=datetime.utcnow, nullable=False)
    last_login = db.Column(db.DateTime, nullable=True)
    
    # Email verification fields
    email_verified = db.Column(db.Boolean, default=False, nullable=False)
    verification_token = db.Column(db.String(100), nullable=True, index=True)
    verification_token_expires = db.Column(db.DateTime, nullable=True)
    
    # Password reset fields
    reset_token = db.Column(db.String(100), nullable=True, index=True)
    reset_token_expires = db.Column(db.DateTime, nullable=True)
    
    # Admin field
    is_admin = db.Column(db.Boolean, default=False, nullable=False, index=True)
    
    # Relationship to translations
    translations = db.relationship('Translation', backref='user', lazy=True, cascade='all, delete-orphan')
    
    def set_password(self, password):
        """Set password hash - uses pbkdf2:sha256 for macOS compatibility"""
        # macOS LibreSSL doesn't support scrypt
        # Werkzeug 3.0+ tries to use scrypt by default, which fails on macOS
        # We MUST bypass Werkzeug entirely on macOS and use pbkdf2_hmac directly
        import hashlib
        import secrets
        
        # Always check for scrypt availability first
        scrypt_available = hasattr(hashlib, 'scrypt')
        
        # On macOS (no scrypt), ALWAYS use direct pbkdf2_hmac - NEVER call Werkzeug
        if not scrypt_available:
            # macOS LibreSSL: Create pbkdf2 hash directly using hashlib
            # This avoids Werkzeug entirely which would try to use scrypt
            try:
                # Generate salt (32 hex chars = 16 bytes)
                salt = secrets.token_hex(16)
                # Use pbkdf2_hmac with sha256
                iterations = 600000  # Werkzeug's default iterations
                hash_bytes = hashlib.pbkdf2_hmac('sha256', password.encode('utf-8'), salt.encode('utf-8'), iterations)
                hash_hex = hash_bytes.hex()
                # Format: pbkdf2:sha256:iterations$salt$hash (Werkzeug-compatible format)
                self.password_hash = f'pbkdf2:sha256:{iterations}${salt}${hash_hex}'
                return  # Success - exit immediately
            except Exception as e:
                # If direct pbkdf2 fails, use simple sha256 as last resort
                # DO NOT fall back to Werkzeug as it will try scrypt and fail
                app.logger.error(f'Error creating pbkdf2 hash directly: {e}')
                salt = secrets.token_hex(16)
                self.password_hash = f'sha256:{salt}:{hashlib.sha256((password + salt).encode()).hexdigest()}'
                app.logger.warning(f'Using simple sha256 hash as last resort on macOS')
                return  # Exit - never call Werkzeug
        
        # Only use Werkzeug if scrypt IS available (Linux/Windows with OpenSSL)
        # On macOS, we never reach here due to the return statements above
        try:
            # Force pbkdf2:sha256 explicitly
            self.password_hash = generate_password_hash(password, method='pbkdf2:sha256')
            # Verify it's not using scrypt
            if self.password_hash.startswith('scrypt'):
                raise ValueError('Werkzeug used scrypt instead of pbkdf2')
            return
        except Exception as e:
            # Fallback: create pbkdf2 hash manually (shouldn't happen on systems with scrypt)
            app.logger.warning(f'Werkzeug failed, creating pbkdf2 hash manually: {e}')
            salt = secrets.token_hex(16)
            iterations = 600000
            hash_bytes = hashlib.pbkdf2_hmac('sha256', password.encode('utf-8'), salt.encode('utf-8'), iterations)
            hash_hex = hash_bytes.hex()
            self.password_hash = f'pbkdf2:sha256:{iterations}${salt}${hash_hex}'
            return
    
    def check_password(self, password):
        """Check password hash"""
        if not self.password_hash:
            return False
        try:
            return check_password_hash(self.password_hash, password)
        except Exception as e:
            # Fallback for custom hash format (sha256 fallback)
            if self.password_hash.startswith('sha256:'):
                try:
                    parts = self.password_hash.split(':')
                    if len(parts) == 3:
                        salt = parts[1]
                        stored_hash = parts[2]
                        import hashlib
                        computed_hash = hashlib.sha256((password + salt).encode()).hexdigest()
                        return computed_hash == stored_hash
                except Exception:
                    pass
            app.logger.error(f'Error checking password: {e}')
            return False
    
    def generate_verification_token(self):
        """Generate email verification token"""
        import secrets
        from datetime import timedelta
        self.verification_token = secrets.token_urlsafe(32)
        self.verification_token_expires = datetime.utcnow() + timedelta(hours=24)  # 24 hours expiry
        return self.verification_token
    
    def generate_reset_token(self):
        """Generate password reset token"""
        import secrets
        from datetime import timedelta
        self.reset_token = secrets.token_urlsafe(32)
        self.reset_token_expires = datetime.utcnow() + timedelta(hours=1)  # 1 hour expiry
        return self.reset_token
    
    def verify_token(self, token):
        """Verify email verification token"""
        if not self.verification_token or not self.verification_token_expires:
            return False
        if datetime.utcnow() > self.verification_token_expires:
            return False
        return self.verification_token == token
    
    def verify_reset_token(self, token):
        """Verify password reset token"""
        if not self.reset_token or not self.reset_token_expires:
            return False
        if datetime.utcnow() > self.reset_token_expires:
            return False
        return self.reset_token == token
    
    def __repr__(self):
        return f'<User {self.email}>'

class Translation(db.Model):
    """Translation history model"""
    __tablename__ = 'translations'
    
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('users.id', ondelete='CASCADE'), nullable=False, index=True)
    original_text = db.Column(db.Text, nullable=False)
    translated_text = db.Column(db.Text, nullable=False)
    detected_language = db.Column(db.String(50))
    target_language = db.Column(db.String(50))
    timestamp = db.Column(db.DateTime, default=datetime.utcnow, nullable=False)
    is_favorite = db.Column(db.Boolean, default=False, nullable=False)
    
    def to_dict(self):
        """Convert model to dictionary"""
        return {
            'id': self.id,
            'original_text': self.original_text,
            'detected_language': self.detected_language,
            'target_language': self.target_language,
            'translated_text': self.translated_text,
            'timestamp': self.timestamp.isoformat(),
            'is_favorite': self.is_favorite
        }

# User loader for Flask-Login
@login_manager.user_loader
def load_user(user_id):
    """Load user by ID for Flask-Login"""
    user = User.query.get(int(user_id))
    if user:
        # Expire the user object so it's refreshed from the database
        db.session.expire(user)
    return user

# Admin decorator
from functools import wraps

def admin_required(f):
    """Decorator to require admin access"""
    @wraps(f)
    @login_required
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated:
            flash('Please log in to access this page.', 'info')
            return redirect(url_for('login'))
        
        # Reload user from database to ensure we have latest is_admin status
        # The user object in Flask-Login session might be stale
        try:
            # Reload user from database
            user_id = current_user.id
            fresh_user = User.query.get(user_id)
            if fresh_user:
                is_admin = bool(getattr(fresh_user, 'is_admin', False))
                app.logger.info(f'Admin check for {fresh_user.email}: is_admin={is_admin}')
            else:
                is_admin = False
                app.logger.warning(f'User {user_id} not found in database')
        except Exception as e:
            app.logger.error(f'Error checking admin status: {e}')
            is_admin = False
        
        if not is_admin:
            flash('Access denied. Admin privileges required.', 'error')
            app.logger.warning(f'Unauthorized admin access attempt by user: {current_user.email}')
            return redirect(url_for('index'))
        
        return f(*args, **kwargs)
    return decorated_function

# Email sending functions
def send_email_direct(subject, recipients, html_body, text_body):
    """Send email using direct SMTP (fallback when Flask-Mail fails)"""
    import smtplib
    from email.mime.text import MIMEText
    from email.mime.multipart import MIMEMultipart
    
    try:
        mail_server = app.config['MAIL_SERVER']
        mail_port = app.config['MAIL_PORT']
        mail_username = app.config['MAIL_USERNAME']
        mail_password = app.config['MAIL_PASSWORD']
        use_tls = app.config['MAIL_USE_TLS']
        use_ssl = app.config['MAIL_USE_SSL']
        
        if not mail_username or not mail_password:
            raise ValueError('MAIL_USERNAME or MAIL_PASSWORD not configured')
        
        # Create message
        msg = MIMEMultipart('alternative')
        msg['Subject'] = subject
        msg['From'] = app.config.get('MAIL_DEFAULT_SENDER', mail_username)
        msg['To'] = ', '.join(recipients)
        
        msg.attach(MIMEText(text_body, 'plain'))
        msg.attach(MIMEText(html_body, 'html'))
        
        # Connect and send
        if use_ssl:
            server = smtplib.SMTP_SSL(mail_server, mail_port, timeout=10)
        else:
            server = smtplib.SMTP(mail_server, mail_port, timeout=10)
            if use_tls:
                server.starttls()
        
        server.login(mail_username, mail_password)
        server.send_message(msg)
        server.quit()
        
        return True
    except Exception as e:
        app.logger.error(f'Direct SMTP error: {str(e)}')
        raise

def send_verification_email(user):
    """Send email verification email to user"""
    try:
        token = user.generate_verification_token()
        db.session.commit()
        
        verification_url = url_for('verify_email', token=token, _external=True)
        
        html_body = render_template('emails/verify_email.html', 
                                   user=user, 
                                   verification_url=verification_url)
        text_body = render_template('emails/verify_email.txt', 
                                  user=user, 
                                  verification_url=verification_url)
        
        if not app.config['MAIL_SUPPRESS_SEND']:
            # Try Flask-Mail first, fallback to direct SMTP
            try:
                msg = Message(
                    subject='Verify Your Email - Translation App',
                    recipients=[user.email],
                    html=html_body,
                    body=text_body
                )
                with app.app_context():
                    mail.send(msg)
                app.logger.info(f'Verification email sent via Flask-Mail to {user.email}')
            except Exception as flask_mail_error:
                # Fallback to direct SMTP
                app.logger.warning(f'Flask-Mail failed ({str(flask_mail_error)}), using direct SMTP...')
                send_email_direct(
                    subject='Verify Your Email - Translation App',
                    recipients=[user.email],
                    html_body=html_body,
                    text_body=text_body
                )
                app.logger.info(f'Verification email sent via direct SMTP to {user.email}')
        else:
            app.logger.info(f'[MAIL SUPPRESSED] Verification email would be sent to {user.email}')
            app.logger.info(f'[MAIL SUPPRESSED] Verification URL: {verification_url}')
        
        return True
    except Exception as e:
        app.logger.error(f'Error sending verification email: {str(e)}', exc_info=True)
        db.session.rollback()
        return False

def send_password_reset_email(user):
    """Send password reset email to user"""
    try:
        token = user.generate_reset_token()
        db.session.commit()
        
        reset_url = url_for('reset_password', token=token, _external=True)
        
        html_body = render_template('emails/reset_password.html', 
                                   user=user, 
                                   reset_url=reset_url)
        text_body = render_template('emails/reset_password.txt', 
                                  user=user, 
                                  reset_url=reset_url)
        
        if not app.config['MAIL_SUPPRESS_SEND']:
            # Try Flask-Mail first, fallback to direct SMTP
            try:
                msg = Message(
                    subject='Reset Your Password - Translation App',
                    recipients=[user.email],
                    html=html_body,
                    body=text_body
                )
                with app.app_context():
                    mail.send(msg)
                app.logger.info(f'Password reset email sent via Flask-Mail to {user.email}')
            except Exception as flask_mail_error:
                # Fallback to direct SMTP
                app.logger.warning(f'Flask-Mail failed ({str(flask_mail_error)}), using direct SMTP...')
                send_email_direct(
                    subject='Reset Your Password - Translation App',
                    recipients=[user.email],
                    html_body=html_body,
                    text_body=text_body
                )
                app.logger.info(f'Password reset email sent via direct SMTP to {user.email}')
        else:
            app.logger.info(f'[MAIL SUPPRESSED] Password reset email would be sent to {user.email}')
            app.logger.info(f'[MAIL SUPPRESSED] Reset URL: {reset_url}')
        
        return True
    except Exception as e:
        app.logger.error(f'Error sending password reset email: {str(e)}', exc_info=True)
        db.session.rollback()
        return False

# Context processor to make variables available to all templates
@app.context_processor
def inject_google_oauth():
    """Make Google OAuth status available to all templates"""
    return {
        'google_oauth_enabled': bool(os.getenv('GOOGLE_CLIENT_ID') and os.getenv('GOOGLE_CLIENT_SECRET'))
    }

# Supported languages mapping
LANGUAGES = {
    'en': 'English',
    'es': 'Spanish',
    'fr': 'French',
    'it': 'Italian',
    'ar': 'Arabic',
    'de': 'German',
    'pt': 'Portuguese',
    'ja': 'Japanese',
    'zh': 'Chinese',
    'ru': 'Russian',
    'ko': 'Korean',
    'hi': 'Hindi',
    'nl': 'Dutch',
    'tr': 'Turkish',
    'pl': 'Polish',
    'sv': 'Swedish',
    'da': 'Danish',
    'no': 'Norwegian',
    'fi': 'Finnish'
}

# Language code mapping for translation
LANG_CODES = {
    'English': 'en',
    'Spanish': 'es',
    'French': 'fr',
    'Italian': 'it',
    'Arabic': 'ar',
    'German': 'de',
    'Portuguese': 'pt',
    'Japanese': 'ja',
    'Chinese': 'zh',
    'Russian': 'ru',
    'Korean': 'ko',
    'Hindi': 'hi',
    'Dutch': 'nl',
    'Turkish': 'tr',
    'Polish': 'pl',
    'Swedish': 'sv',
    'Danish': 'da',
    'Norwegian': 'no',
    'Finnish': 'fi'
}

def detect_language(text):
    """Detect the language of the input text"""
    try:
        lang_code = detect(text)
        return LANGUAGES.get(lang_code, 'Unknown')
    except:
        return 'Unknown'

def translate_text(text, target_lang):
    """Translate text to target language"""
    try:
        # Detect source language
        source_lang = detect(text)
        target_code = LANG_CODES[target_lang]
        
        # If source and target are the same, return original
        if source_lang == target_code:
            return text
        
        # Translate using Google Translator
        translator = GoogleTranslator(source=source_lang, target=target_code)
        translated = translator.translate(text)
        return translated
    except Exception as e:
        return f"Translation error: {str(e)}"
    
def detect_intent(text: str) -> str:
    """
    Phase 1 Intent Detector (Rule-based)
    Returns one of:
    translation | writing | meeting | content | general_assistant
    """
    if not text:
        return "general_assistant"

    t = text.strip().lower()

    # Arabic + English keywords
    if any(k in t for k in ["ترجم", "ترجمة", "translate", "translation"]):
        return "translation"

    if any(k in t for k in ["لخص", "تلخيص", "summarize", "summary"]):
        return "writing"

    if any(k in t for k in ["اجتماع", "meeting", "minutes", "محضر"]):
        return "meeting"

    if any(k in t for k in ["قصة", "story", "اكتب", "write a", "generate content", "content"]):
        return "content"

    return "general_assistant"

def orchestrator_handle(intent: str, text: str, meta: dict | None = None):
    """
    Phase 1 Orchestrator:
    - translation: uses existing translate_text()
    - others: placeholders (Phase 2+)
    """
    meta = meta or {}

    if intent == "translation":
        target_lang = meta.get("target_lang", "English")  # default
        translated_text = translate_text(text, target_lang)
        return {
            "type": "translation",
            "target_language": target_lang,
            "translated_text": translated_text
        }

    if intent == "writing":
        # Phase 2 will implement writing assistant
        return {
            "type": "writing",
            "message": "Writing Assistant not implemented yet (Phase 2).",
            "hint": "Try: 'لخص النص التالي ...' في Phase 2"
        }

    if intent == "meeting":
        # Phase 4 will implement meeting assistant
        return {
            "type": "meeting",
            "message": "Meeting Assistant not implemented yet (Phase 4).",
            "hint": "Try: 'لخص الاجتماع واستخرج المهام' في Phase 4"
        }

    if intent == "content":
        # Phase 5 will implement content generator
        return {
            "type": "content",
            "message": "Content Generator not implemented yet (Phase 5).",
            "hint": "Try: 'اكتب قصة قصيرة عن ...' في Phase 5"
        }

    # general assistant fallback
    return {
        "type": "general_assistant",
        "message": "I detected a general request. In Phase 1 we focus on routing + translation."
    }



# Error handler for rate limit exceeded
@app.errorhandler(RateLimitExceeded)
def handle_rate_limit(e):
    """Handle rate limit exceeded errors with user-friendly messages"""
    client_ip = get_remote_address()
    app.logger.warning(f'Rate limit exceeded for IP: {client_ip}')
    return jsonify({
        'error': 'Rate limit exceeded',
        'message': 'You have made too many requests. Please wait a moment before trying again.',
        'retry_after': getattr(e, 'retry_after', None)
    }), 429

# Global error handlers
@app.errorhandler(404)
def handle_404(e):
    """Handle 404 Not Found errors"""
    app.logger.warning(f'404 Not Found: {request.path} from {get_remote_address()}')
    if request.path.startswith('/api/') or request.is_json:
        return jsonify({'error': 'Not found', 'message': 'The requested resource was not found'}), 404
    # Redirect to login if not authenticated, otherwise show 404
    if not current_user.is_authenticated:
        return redirect(url_for('login'))
    return render_template('index.html', user=current_user), 404

@app.errorhandler(500)
def handle_500(e):
    """Handle 500 Internal Server errors"""
    client_ip = get_remote_address()
    app.logger.error(f'500 Internal Server Error from {client_ip}: {str(e)}', exc_info=True)
    
    # Capture exception in Sentry if enabled
    if sentry_dsn:
        try:
            import sentry_sdk
            sentry_sdk.capture_exception(e)
        except Exception:
            pass
    
    if request.path.startswith('/api/') or request.is_json:
        return jsonify({
            'error': 'Internal server error',
            'message': 'An unexpected error occurred. Please try again later.'
        }), 500
    
    # Always redirect to login for unauthenticated users or any errors
    # Never try to render index.html in error handlers as it requires user context
    try:
        # Try to check if user is authenticated
        if hasattr(current_user, 'is_authenticated') and not current_user.is_authenticated:
            flash('An error occurred. Please try again.', 'error')
            return redirect(url_for('login'))
    except Exception:
        # If checking auth fails, still redirect to login
        pass
    
    # For any error, redirect to login (safer than trying to render templates)
    flash('An error occurred. Please try again.', 'error')
    return redirect(url_for('login'))

@app.errorhandler(Exception)
def handle_exception(e):
    """Handle all unhandled exceptions"""
    client_ip = get_remote_address()
    app.logger.error(f'Unhandled exception from {client_ip}: {str(e)}', exc_info=True)
    
    # Capture exception in Sentry if enabled
    if sentry_dsn:
        try:
            import sentry_sdk
            sentry_sdk.capture_exception(e)
        except Exception:
            pass
    
    # Don't handle HTTP exceptions (let Flask handle them)
    from werkzeug.exceptions import HTTPException
    if isinstance(e, HTTPException):
        return e
    
    if request.path.startswith('/api/') or request.is_json:
        return jsonify({
            'error': 'Internal server error',
            'message': 'An unexpected error occurred. Please try again later.'
        }), 500
    
    # Always redirect to login for unauthenticated users or any errors
    # Never try to render index.html in error handlers as it requires user context
    try:
        # Try to check if user is authenticated
        if hasattr(current_user, 'is_authenticated') and not current_user.is_authenticated:
            flash('An error occurred. Please try again.', 'error')
            return redirect(url_for('login'))
    except Exception:
        # If checking auth fails, still redirect to login
        pass
    
    # For any error, redirect to login (safer than trying to render templates)
    flash('An error occurred. Please try again.', 'error')
    return redirect(url_for('login'))

@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint for deployment platforms and monitoring"""
    health_status = {
        'status': 'healthy',
        'timestamp': datetime.now().isoformat(),
        'version': '1.0.0',
        'checks': {}
    }
    
    # Check database connectivity
    db_status = 'healthy'
    db_error = None
    try:
        # Simple query to test database connection
        from sqlalchemy import text
        db.session.execute(text('SELECT 1'))
        db.session.commit()
        health_status['checks']['database'] = {
            'status': 'healthy',
            'message': 'Database connection successful'
        }
    except Exception as e:
        db_status = 'unhealthy'
        db_error = str(e)
        health_status['checks']['database'] = {
            'status': 'unhealthy',
            'message': f'Database connection failed: {db_error}'
        }
        app.logger.warning(f'Health check: Database connection failed - {db_error}')
    
    # Check Sentry status
    sentry_status = 'enabled' if sentry_dsn else 'disabled'
    health_status['checks']['sentry'] = {
        'status': sentry_status,
        'message': 'Sentry error tracking enabled' if sentry_dsn else 'Sentry error tracking disabled'
    }
    
    # Overall status
    if db_status == 'unhealthy':
        health_status['status'] = 'unhealthy'
        return jsonify(health_status), 503  # Service Unavailable
    
    return jsonify(health_status), 200

# Authentication Routes
@app.route('/login', methods=['GET', 'POST'])
@limiter.limit("10 per minute")
def login():
    """Login page and handler"""
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    
    if request.method == 'POST':
        email = request.form.get('email', '').strip().lower()
        password = request.form.get('password', '')
        remember = bool(request.form.get('remember'))
        
        if not email or not password:
            flash('Please provide both email and password.', 'error')
            return render_template('login.html'), 400
        
        # Find user by email
        user = User.query.filter_by(email=email, provider='local').first()
        
        if user and user.check_password(password):
            # Check if email is verified
            if not user.email_verified:
                flash('Please verify your email address before logging in. Check your inbox for the verification link.', 'warning')
                return render_template('login.html', show_resend_verification=True, user_email=email), 401
            
            user.last_login = datetime.utcnow()
            db.session.commit()
            login_user(user, remember=remember)
            app.logger.info(f'User logged in: {user.email}')
            
            # Handle redirect after login
            # Flask-Login stores next URL in session or request.args
            from urllib.parse import urlparse
            next_page = request.args.get('next')
            
            # If not in args, check session (Flask-Login may store it there)
            if not next_page:
                next_page = session.pop('next', None) or session.pop('_next', None)
            
            # Validate and redirect
            if next_page:
                parsed = urlparse(next_page)
                # Only allow relative URLs (no netloc) or same host
                # Also allow paths starting with / (safe internal redirects)
                if not parsed.netloc or parsed.netloc == request.host.split(':')[0]:
                    app.logger.info(f'Login redirect: {next_page}')
                    return redirect(next_page)
                else:
                    app.logger.warning(f'Invalid redirect attempt to: {next_page}')
            
            app.logger.info('Login redirect: index (no next_page)')
            return redirect(url_for('index'))
        else:
            flash('Invalid email or password.', 'error')
            app.logger.warning(f'Failed login attempt for email: {email}')
            return render_template('login.html'), 401
    
    return render_template('login.html')

@app.route('/signup', methods=['GET', 'POST'])
@limiter.limit("5 per minute")
def signup():
    """Signup page and handler"""
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    
    if request.method == 'POST':
        email = request.form.get('email', '').strip().lower()
        username = request.form.get('username', '').strip()
        password = request.form.get('password', '')
        confirm_password = request.form.get('confirm_password', '')
        name = request.form.get('name', '').strip()
        
        # Validation
        if not email or not password:
            flash('Email and password are required.', 'error')
            return render_template('signup.html'), 400
        
        if password != confirm_password:
            flash('Passwords do not match.', 'error')
            return render_template('signup.html'), 400
        
        if len(password) < 8:
            flash('Password must be at least 8 characters long.', 'error')
            return render_template('signup.html'), 400
        
        # Check if user already exists
        if User.query.filter_by(email=email).first():
            flash('Email already registered. Please login instead.', 'error')
            return render_template('signup.html'), 409
        
        if username and User.query.filter_by(username=username).first():
            flash('Username already taken.', 'error')
            return render_template('signup.html'), 409
        
        # Create new user
        user = User(
            email=email,
            username=username or None,
            name=name or email.split('@')[0],
            provider='local',
            email_verified=False  # Email not verified yet
        )
        user.set_password(password)
        
        try:
            db.session.add(user)
            db.session.commit()
            app.logger.info(f'New user registered: {user.email}')
            
            # Send verification email
            email_sent = send_verification_email(user)
            if email_sent:
                flash('Account created successfully! Please check your email to verify your account.', 'success')
                app.logger.info(f'Verification email sent to {user.email}')
            else:
                flash('Account created, but we couldn\'t send the verification email. Please check your email configuration or contact support.', 'warning')
                app.logger.warning(f'Failed to send verification email to {user.email} - check email configuration')
            
            # Don't auto-login - require email verification first
            return redirect(url_for('login'))
        except Exception as e:
            db.session.rollback()
            app.logger.error(f'Error creating user: {str(e)}', exc_info=True)
            flash('An error occurred. Please try again.', 'error')
            return render_template('signup.html'), 500
    
    return render_template('signup.html')

@app.route('/logout')
@login_required
def logout():
    """Logout handler"""
    email = current_user.email
    logout_user()
    app.logger.info(f'User logged out: {email}')
    flash('You have been logged out.', 'info')
    return redirect(url_for('login'))

@app.route('/verify-email/<token>')
@limiter.limit("10 per hour")
def verify_email(token):
    """Verify user email with token"""
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    
    user = User.query.filter_by(verification_token=token).first()
    
    if not user:
        flash('Invalid or expired verification link.', 'error')
        return redirect(url_for('login'))
    
    if user.verify_token(token):
        user.email_verified = True
        user.verification_token = None
        user.verification_token_expires = None
        db.session.commit()
        app.logger.info(f'Email verified for user: {user.email}')
        flash('Email verified successfully! You can now log in.', 'success')
        return redirect(url_for('login'))
    else:
        flash('Verification link has expired. Please request a new one.', 'error')
        return redirect(url_for('resend_verification', email=user.email))

@app.route('/resend-verification', methods=['GET', 'POST'])
@limiter.limit("5 per hour")
def resend_verification():
    """Resend email verification"""
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    
    if request.method == 'POST':
        email = request.form.get('email', '').strip().lower()
        
        if not email:
            flash('Please provide your email address.', 'error')
            return render_template('resend_verification.html')
        
        user = User.query.filter_by(email=email, provider='local').first()
        
        if not user:
            # Don't reveal if email exists for security
            flash('If an account exists with this email, a verification link has been sent.', 'info')
            return redirect(url_for('login'))
        
        if user.email_verified:
            flash('Your email is already verified. You can log in.', 'info')
            return redirect(url_for('login'))
        
        if send_verification_email(user):
            flash('Verification email sent! Please check your inbox.', 'success')
        else:
            flash('Failed to send verification email. Please try again later.', 'error')
        
        return redirect(url_for('login'))
    
    # GET request - show form
    email = request.args.get('email', '')
    return render_template('resend_verification.html', email=email)

@app.route('/forgot-password', methods=['GET', 'POST'])
@limiter.limit("5 per hour")
def forgot_password():
    """Forgot password - send reset link"""
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    
    if request.method == 'POST':
        email = request.form.get('email', '').strip().lower()
        
        if not email:
            flash('Please provide your email address.', 'error')
            return render_template('forgot_password.html')
        
        user = User.query.filter_by(email=email, provider='local').first()
        
        if not user:
            # Don't reveal if email exists for security
            flash('If an account exists with this email, a password reset link has been sent.', 'info')
            return redirect(url_for('login'))
        
        if send_password_reset_email(user):
            flash('Password reset link sent! Please check your email.', 'success')
        else:
            flash('Failed to send password reset email. Please try again later.', 'error')
        
        return redirect(url_for('login'))
    
    return render_template('forgot_password.html')

@app.route('/reset-password/<token>', methods=['GET', 'POST'])
@limiter.limit("10 per hour")
def reset_password(token):
    """Reset password with token"""
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    
    user = User.query.filter_by(reset_token=token).first()
    
    if not user:
        flash('Invalid or expired password reset link.', 'error')
        return redirect(url_for('forgot_password'))
    
    if not user.verify_reset_token(token):
        flash('Password reset link has expired. Please request a new one.', 'error')
        return redirect(url_for('forgot_password'))
    
    if request.method == 'POST':
        password = request.form.get('password', '')
        confirm_password = request.form.get('confirm_password', '')
        
        if not password:
            flash('Password is required.', 'error')
            return render_template('reset_password.html', token=token)
        
        if password != confirm_password:
            flash('Passwords do not match.', 'error')
            return render_template('reset_password.html', token=token)
        
        if len(password) < 8:
            flash('Password must be at least 8 characters long.', 'error')
            return render_template('reset_password.html', token=token)
        
        # Reset password
        user.set_password(password)
        user.reset_token = None
        user.reset_token_expires = None
        db.session.commit()
        
        app.logger.info(f'Password reset for user: {user.email}')
        flash('Password reset successfully! You can now log in with your new password.', 'success')
        return redirect(url_for('login'))
    
    return render_template('reset_password.html', token=token)

@app.route('/login/google')
@limiter.limit("10 per minute")
def login_google():
    """Initiate Google OAuth login"""
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    
    # Check if Google OAuth is configured
    if not os.getenv('GOOGLE_CLIENT_ID') or not os.getenv('GOOGLE_CLIENT_SECRET'):
        flash('Google login is not configured. Please contact administrator.', 'error')
        return redirect(url_for('login'))
    
    redirect_uri = url_for('auth_google_callback', _external=True)
    # Preserve next parameter if it exists
    state = request.args.get('next', '')
    return google.authorize_redirect(redirect_uri, state=state)

@app.route('/auth/google/callback')
@limiter.limit("10 per minute")
def auth_google_callback():
    """Handle Google OAuth callback"""
    try:
        token = google.authorize_access_token()
        user_info = token.get('userinfo')
        
        if not user_info:
            flash('Failed to retrieve user information from Google.', 'error')
            return redirect(url_for('login'))
        
        email = user_info.get('email', '').lower()
        name = user_info.get('name', '')
        picture = user_info.get('picture', '')
        provider_id = user_info.get('sub', '')
        
        if not email:
            flash('Email not provided by Google.', 'error')
            return redirect(url_for('login'))
        
        # Find or create user
        user = User.query.filter_by(email=email, provider='google').first()
        
        if not user:
            # Check if email exists with different provider
            existing_user = User.query.filter_by(email=email).first()
            if existing_user:
                flash('This email is already registered with a different login method.', 'error')
                return redirect(url_for('login'))
            
            # Create new user
            user = User(
                email=email,
                name=name or email.split('@')[0],
                picture=picture,
                provider='google',
                provider_id=provider_id,
                email_verified=True  # Google OAuth users are automatically verified
            )
            db.session.add(user)
            db.session.commit()
            app.logger.info(f'New Google user registered: {user.email}')
        else:
            # Update user info
            user.name = name or user.name
            user.picture = picture or user.picture
            user.last_login = datetime.utcnow()
            user.email_verified = True  # Google OAuth users are automatically verified
            db.session.commit()
            app.logger.info(f'Google user logged in: {user.email}')
        
        login_user(user, remember=True)
        flash('Logged in successfully with Google!', 'success')
        
        # Handle redirect after login (check session for next_page set by Flask-Login)
        next_page = request.args.get('next') or session.get('next')
        if next_page:
            from urllib.parse import urlparse
            parsed = urlparse(next_page)
            if parsed.netloc == '' or parsed.netloc == request.host:
                session.pop('next', None)
                return redirect(next_page)
        return redirect(url_for('index'))
    
    except Exception as e:
        app.logger.error(f'Google OAuth error: {str(e)}', exc_info=True)
        flash('An error occurred during Google login. Please try again.', 'error')
        return redirect(url_for('login'))

@app.route('/')
@login_required
def index():
    """Main application page - requires authentication"""
    # Refresh user from database to get latest data (like is_admin)
    # This ensures we have the most up-to-date user object
    try:
        db.session.refresh(current_user)
    except Exception:
        # If refresh fails, reload the user
        user_id = current_user.id
        user = User.query.get(user_id)
        if user:
            login_user(user, remember=True)
    return render_template('index.html', user=current_user)

@app.route('/test-error')
def test_error():
    """Test route to verify Sentry error tracking"""
    # Allow test endpoint in all environments (can be disabled via env var)
    if os.getenv('DISABLE_TEST_ENDPOINT', 'false').lower() == 'true':
        return jsonify({'error': 'Test endpoint disabled'}), 403
    
    # Check if Sentry is enabled - reload from env to ensure we have latest value
    from dotenv import load_dotenv
    load_dotenv()  # Reload to ensure .env is read
    current_sentry_dsn = os.getenv('SENTRY_DSN')
    sentry_enabled = bool(current_sentry_dsn)
    
    # Also check module-level variable for comparison
    module_dsn = sentry_dsn if 'sentry_dsn' in globals() else None
    
    # Capture a test exception
    try:
        raise Exception("This is a test error for Sentry error tracking")
    except Exception as e:
        app.logger.error("Test error triggered", exc_info=True)
        if sentry_enabled:
            try:
                import sentry_sdk
                sentry_sdk.capture_exception(e)
                app.logger.info("Test error sent to Sentry")
            except Exception as sentry_error:
                app.logger.warning(f"Failed to send to Sentry: {sentry_error}")
        
        return jsonify({
            'message': 'Test error sent to Sentry (if configured)',
            'sentry_enabled': sentry_enabled,
            'sentry_dsn_from_env': bool(current_sentry_dsn),
            'module_level_dsn': bool(module_dsn),
            'debug': {
                'env_dsn_length': len(current_sentry_dsn) if current_sentry_dsn else 0,
                'module_dsn_length': len(module_dsn) if module_dsn else 0
            }
        }), 500

@app.route('/translate', methods=['POST'])
@login_required
@limiter.limit("10 per minute")  # Allow 10 translations per minute per IP
def translate():
    client_ip = get_remote_address()
    try:
        data = request.get_json()
        text = data.get('text', '').strip()
        target_lang = data.get('target_lang', 'English')
        
        app.logger.info(f'Translation request from {current_user.email} ({client_ip}): {len(text)} chars, target: {target_lang}')
        
        if not text:
            app.logger.warning(f'Empty translation request from {current_user.email} ({client_ip})')
            # Optionally send validation errors to Sentry (can be disabled via env var)
            capture_validation = os.getenv('SENTRY_CAPTURE_VALIDATION_ERRORS', 'false').lower() == 'true'
            # Check both module-level variable and environment variable
            current_sentry_dsn = sentry_dsn if 'sentry_dsn' in globals() and sentry_dsn else os.getenv('SENTRY_DSN')
            
            if current_sentry_dsn and capture_validation:
                try:
                    import sentry_sdk
                    sentry_sdk.capture_message(
                        f'Empty translation request from {current_user.email} ({client_ip})',
                        level='warning'
                    )
                    app.logger.info(f'✅ Empty translation request sent to Sentry from {current_user.email}')
                except Exception as e:
                    app.logger.warning(f'Failed to send validation error to Sentry: {e}')
            elif capture_validation and not current_sentry_dsn:
                app.logger.debug('Validation error tracking enabled but Sentry DSN not set')
            elif not capture_validation:
                app.logger.debug('Validation error tracking is disabled (set SENTRY_CAPTURE_VALIDATION_ERRORS=true to enable)')
            return jsonify({'error': 'No text provided'}), 400
        
        # Detect source language
        detected_lang = detect_language(text)
        app.logger.debug(f'Detected language: {detected_lang} for text: {text[:50]}...')
        
        # Translate text
        translated_text = translate_text(text, target_lang)
        
        # Save to database with user_id
        translation = Translation(
            user_id=current_user.id,
            original_text=text,
            detected_language=detected_lang,
            target_language=target_lang,
            translated_text=translated_text,
            is_favorite=False
        )
        db.session.add(translation)
        db.session.commit()
        app.logger.info(f'Translation saved to database (ID: {translation.id}, User: {current_user.email})')
        
        # Keep only last 100 translations per user (optional cleanup)
        user_translations_count = Translation.query.filter_by(user_id=current_user.id).count()
        if user_translations_count > 100:
            # Delete oldest translations beyond 100 for this user
            oldest = Translation.query.filter_by(user_id=current_user.id).order_by(Translation.timestamp.asc()).limit(user_translations_count - 100).all()
            deleted_count = len(oldest)
            for old_translation in oldest:
                db.session.delete(old_translation)
            db.session.commit()
            app.logger.info(f'Cleaned up {deleted_count} old translations for user {current_user.email} (kept last 100)')
        
        app.logger.info(f'Translation successful: {detected_lang} -> {target_lang}')
        return jsonify({
            'original_text': text,
            'detected_language': detected_lang,
            'target_language': target_lang,
            'translated_text': translated_text,
            'history_id': translation.id
        })
    except Exception as e:
        app.logger.error(f'Translation error from {current_user.email} ({client_ip}): {str(e)}', exc_info=True)
        return jsonify({'error': str(e)}), 500

ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx'}

def allowed_file(filename):
    """التحقق من الامتداد المسموح به للملف"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/upload_text_file', methods=['POST'])
@login_required
@limiter.limit("5 per minute")  # Allow 5 file uploads per minute per IP
def upload_text_file():
    try:
        # التحقق من وجود الملف في الطلب
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        
        # التحقق من اسم الملف
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        # التحقق من امتداد الملف
        if not allowed_file(file.filename):
            return jsonify({'error': 'Invalid file type. Please upload a .txt, .pdf, or .docx file'}), 400
        
        # حفظ الملف باستخدام اسم آمن
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        # التعامل مع أنواع الملفات المختلفة
        if filename.endswith('.txt'):
            with open(filepath, 'r', encoding='utf-8') as f:
                text = f.read()
        elif filename.endswith('.pdf'):
            from PyPDF2 import PdfReader
            text = ''
            with open(filepath, 'rb') as f:
                reader = PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text()
        elif filename.endswith('.docx'):
            from docx import Document
            text = ''
            doc = Document(filepath)
            for para in doc.paragraphs:
                text += para.text
        
        # ترجم النص
        target_lang = request.form.get('target_lang', 'English')
        detected_lang = detect_language(text)
        translated_text = translate_text(text, target_lang)

        # حفظ الترجمة في قاعدة البيانات
        translation = Translation(
            user_id=current_user.id,
            original_text=text,
            detected_language=detected_lang,
            target_language=target_lang,
            translated_text=translated_text,
            is_favorite=False
        )
        db.session.add(translation)
        db.session.commit()

        return jsonify({
            'original_text': text,
            'detected_language': detected_lang,
            'target_language': target_lang,
            'translated_text': translated_text,
            'history_id': translation.id
        })

    except Exception as e:
        app.logger.error(f'Error uploading file: {str(e)}', exc_info=True)
        return jsonify({'error': 'Internal server error'}), 500
    
@app.route('/summarize_text', methods=['POST'])
def summarize_text():
    try:
        data = request.get_json()
        text = data.get('text', '')

        if not text:
            return jsonify({'error': 'No text provided for summarization'}), 400
        word_count = len(text.split())
        if word_count < 40:
            return jsonify({'summarized_text': text})

        # Abstractive summary (rephrased, shorter) with timeout fallback
        summarizer_model = get_summarizer()
        trimmed_text = text[:800]
        prompt = f"summarize: {trimmed_text}"

        def _run_summary():
            return summarizer_model(prompt, max_length=70, min_length=25, do_sample=False)

        with concurrent.futures.ThreadPoolExecutor(max_workers=1) as executor:
            future = executor.submit(_run_summary)
            try:
                summarized = future.result(timeout=5)
                return jsonify({'summarized_text': summarized[0]['generated_text']})
            except Exception:
                # Fallback to fast extractive summary if model is slow or fails
                return jsonify({'summarized_text': quick_summary(text)})

    except Exception as e:
        app.logger.error(f"Error in summarization: {str(e)}", exc_info=True)
        return jsonify({'error': 'Summarization failed'}), 500


@app.route('/upload_audio', methods=['POST'])
@login_required
@limiter.limit("5 per minute")
def upload_audio():
    client_ip = get_remote_address()
    filepath = None
    app.logger.info(f"request.files: {request.files}")
    app.logger.info(f"request.form: {request.form}")

    try:
        if 'audio' not in request.files:
            return jsonify({'error': 'No audio file provided'}), 400

        audio_file = request.files['audio']
        target_lang = request.form.get('target_lang', 'English')

        if audio_file.filename == '':
            return jsonify({'error': 'Empty audio file'}), 400

        # 1️⃣ حفظ الملف باسم مؤقت فريد (حل WinError)
        filename = f"{uuid.uuid4().hex}.webm"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        audio_file.save(filepath)

        if not os.path.exists(filepath):
            raise FileNotFoundError("Audio file was not saved correctly")

        app.logger.info(f"Audio saved: {filepath}")

        # 2️⃣ تحويل الصوت إلى نص (Whisper)
        text = transcribe_audio(filepath)

        if not text:
            return jsonify({'error': 'No speech detected'}), 400

        # 3️⃣ كشف اللغة والترجمة
        detected_lang = detect_language(text)
        translated_text = translate_text(text, target_lang)

        # 4️⃣ حفظ النتيجة في قاعدة البيانات
        translation = Translation(
            user_id=current_user.id,
            original_text=text,
            detected_language=detected_lang,
            target_language=target_lang,
            translated_text=translated_text,
            is_favorite=False
        )
        db.session.add(translation)
        db.session.commit()

        return jsonify({
            'original_text': text,
            'detected_language': detected_lang,
            'target_language': target_lang,
            'translated_text': translated_text,
            'history_id': translation.id
        })

    except Exception as e:
        app.logger.error(f"Audio upload error from {client_ip}: {e}", exc_info=True)
        return jsonify({'error': 'Audio processing failed'}), 500

    finally:
        # 5️⃣ حذف الملف بعد انتهاء Whisper (Windows-safe)
        if filepath and os.path.exists(filepath):
            try:
                time.sleep(1)  # مهم جدًا في Windows
                os.remove(filepath)
            except PermissionError:
                app.logger.warning(f"File still in use: {filepath}")
    client_ip = get_remote_address()

    if 'audio' not in request.files:
        return jsonify({'error': 'No audio file provided'}), 400

    audio_file = request.files['audio']
    target_lang = request.form.get('target_lang', 'English')

    if audio_file.filename == '':
        return jsonify({'error': 'Empty filename'}), 400

    # ✅ أنشئ اسم ملف آمن ومؤقت
    filename = f"{uuid.uuid4()}.webm"
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)

    try:
        # 1️⃣ احفظ الملف
        audio_file.save(filepath)

        # 🔍 تأكيد وجوده (مهم للتشخيص)
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"Saved file not found: {filepath}")

        app.logger.info(f"Audio saved at {filepath}")

        # 2️⃣ حوّل الصوت إلى نص (Whisper)
        text = transcribe_audio(filepath)

        if text.startswith("Error"):
            return jsonify({'error': text}), 400

        # 3️⃣ اكتشاف اللغة + الترجمة
        detected_lang = detect_language(text)
        translated_text = translate_text(text, target_lang)

        # 4️⃣ الحفظ في DB
        translation = Translation(
            user_id=current_user.id,
            original_text=text,
            detected_language=detected_lang,
            target_language=target_lang,
            translated_text=translated_text
        )
        db.session.add(translation)
        db.session.commit()

        return jsonify({
            'original_text': text,
            'translated_text': translated_text,
            'detected_language': detected_lang
        })

    except Exception as e:
        app.logger.error(f"Audio upload error: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500

    finally:
        # 5️⃣ احذف الملف بعد انتهاء Whisper
        if os.path.exists(filepath):
            try:
                os.remove(filepath)
            except Exception as e:
                app.logger.warning(f"Failed to delete temp file: {e}")
    client_ip = get_remote_address()
    filepath = None

    try:
        # 1️⃣ التحقق من وجود الملف
        if 'audio' not in request.files:
            app.logger.warning(f'Audio upload: No file provided from {client_ip}')
            return jsonify({'error': 'No audio file provided'}), 400

        audio_file = request.files['audio']
        target_lang = request.form.get('target_lang', 'English')

        if not audio_file or audio_file.filename == '':
            app.logger.warning(f'Audio upload: Empty filename from {client_ip}')
            return jsonify({'error': 'No audio file selected'}), 400

        app.logger.info(
            f'Audio upload from {client_ip}: {audio_file.filename}, target={target_lang}'
        )

        # 2️⃣ إنشاء اسم ملف مؤقت فريد (حل WinError 32)
        filename = f"{uuid.uuid4().hex}_{secure_filename(audio_file.filename)}"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)

        audio_file.save(filepath)
        file_size = os.path.getsize(filepath)

        app.logger.debug(
            f'Audio file saved: {filepath}, size={file_size} bytes'
        )

        # 3️⃣ تحويل الصوت إلى نص باستخدام Whisper
        text = transcribe_audio(filepath)

        if not text or text.lower().startswith('error'):
            app.logger.error(f'Audio transcription failed: {text}')
            return jsonify({'error': text or 'Audio transcription failed'}), 400

        app.logger.info(
            f'Audio transcribed successfully ({len(text)} characters)'
        )

        # 4️⃣ كشف اللغة
        detected_lang = detect_language(text)
        app.logger.debug(f'Detected language: {detected_lang}')

        # 5️⃣ الترجمة
        translated_text = translate_text(text, target_lang)

        # 6️⃣ الحفظ في قاعدة البيانات
        translation = Translation(
            user_id=current_user.id,
            original_text=text,
            detected_language=detected_lang,
            target_language=target_lang,
            translated_text=translated_text,
            is_favorite=False
        )

        db.session.add(translation)
        db.session.commit()

        app.logger.info(
            f'Audio translation saved '
            f'(ID={translation.id}, User={current_user.email})'
        )

        # 7️⃣ تنظيف السجل (الاحتفاظ بآخر 100 ترجمة)
        translations_count = Translation.query.filter_by(
            user_id=current_user.id
        ).count()

        if translations_count > 100:
            excess = translations_count - 100
            old_items = (
                Translation.query
                .filter_by(user_id=current_user.id)
                .order_by(Translation.timestamp.asc())
                .limit(excess)
                .all()
            )

            for item in old_items:
                db.session.delete(item)

            db.session.commit()
            app.logger.info(
                f'Cleaned up {len(old_items)} old translations '
                f'for user {current_user.email}'
            )

        # 8️⃣ الاستجابة النهائية
        return jsonify({
            'original_text': text,
            'detected_language': detected_lang,
            'target_language': target_lang,
            'translated_text': translated_text,
            'history_id': translation.id
        })

    except Exception as e:
        app.logger.error(
            f'Audio upload error from {client_ip}: {str(e)}',
            exc_info=True
        )
        return jsonify({'error': 'Internal server error'}), 500

    finally:
        # 9️⃣ تنظيف الملف المؤقت (Windows-safe)
        if filepath and os.path.exists(filepath):
            try:
                time.sleep(0.7)  # مهم لتحرير الملف من Whisper / ffmpeg
                os.remove(filepath)
                app.logger.debug(
                    f'Temporary audio file deleted: {filepath}'
                )
            except PermissionError:
                app.logger.warning(
                    f'File still in use, could not delete: {filepath}'
                )
    client_ip = get_remote_address()
    try:
        if 'audio' not in request.files:
            app.logger.warning(f'Audio upload: No file provided from {client_ip}')
            return jsonify({'error': 'No audio file provided'}), 400
        
        audio_file = request.files['audio']
        target_lang = request.form.get('target_lang', 'English')
        
        if audio_file.filename == '':
            app.logger.warning(f'Audio upload: Empty filename from {client_ip}')
            return jsonify({'error': 'No audio file selected'}), 400
        
        app.logger.info(f'Audio upload from {client_ip}: {audio_file.filename}, target: {target_lang}')
        
        # Save uploaded file temporarily
        filename = secure_filename(audio_file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        audio_file.save(filepath)
        file_size = os.path.getsize(filepath)
        app.logger.debug(f'Audio file saved: {filepath}, size: {file_size} bytes')
        
        try:
            # Transcribe audio to text
            text = transcribe_audio(filepath)
            
            if text.startswith('Error') or text.startswith('Could not'):
                app.logger.error(f'Audio transcription failed: {text}')
                return jsonify({'error': text}), 400
            
            app.logger.info(f'Audio transcribed successfully: {len(text)} chars')
            
            # Detect source language
            detected_lang = detect_language(text)
            app.logger.debug(f'Audio translation: Detected {detected_lang}')
            
            # Translate text
            translated_text = translate_text(text, target_lang)
            
            # Save to database with user_id
            translation = Translation(
                user_id=current_user.id,
                original_text=text,
                detected_language=detected_lang,
                target_language=target_lang,
                translated_text=translated_text,
                is_favorite=False
            )
            db.session.add(translation)
            db.session.commit()
            app.logger.info(f'Audio translation saved (ID: {translation.id}, User: {current_user.email})')
            
            # Keep only last 100 translations per user (optional cleanup)
            user_translations_count = Translation.query.filter_by(user_id=current_user.id).count()
            if user_translations_count > 100:
                oldest = Translation.query.filter_by(user_id=current_user.id).order_by(Translation.timestamp.asc()).limit(user_translations_count - 100).all()
                deleted_count = len(oldest)
                for old_translation in oldest:
                    db.session.delete(old_translation)
                db.session.commit()
                app.logger.info(f'Cleaned up {deleted_count} old translations for user {current_user.email}')
            
            return jsonify({
                'original_text': text,
                'detected_language': detected_lang,
                'target_language': target_lang,
                'translated_text': translated_text,
                'history_id': translation.id
            })
        finally:
            # Clean up temporary file
            if os.path.exists(filepath):
                os.remove(filepath)
                app.logger.debug(f'Cleaned up temporary audio file: {filepath}')
    except Exception as e:
        app.logger.error(f'Audio upload error from {client_ip}: {str(e)}', exc_info=True)
        return jsonify({'error': str(e)}), 500


    """s
    Endpoint to receive audio or video files and process them.
    Extract text (transcript) from the video/audio file and pass it to the translation system.
    """
    try:
        if request.content_type and "multipart/form-data" in request.content_type:
            if "audio" in request.files:
                audio_file = request.files["audio"]
                filename = secure_filename(audio_file.filename)
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                audio_file.save(filepath)

                # Extract text from audio (existing logic)
                transcript = transcribe_audio(filepath)

                if os.path.exists(filepath):
                    os.remove(filepath)

            elif "video" in request.files:
                video_file = request.files["video"]
                filename = secure_filename(video_file.filename)
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                video_file.save(filepath)

                # Extract audio from video using moviepy
                audio_filepath = filepath.replace(".mp4", ".wav")
                video_clip = VideoFileClip(filepath)
                video_clip.audio.write_audiofile(audio_filepath)

                # Extract text from audio (existing logic)
                transcript = transcribe_audio(audio_filepath)

                if os.path.exists(filepath):
                    os.remove(filepath)
                if os.path.exists(audio_filepath):
                    os.remove(audio_filepath)

            else:
                return jsonify({"error": "No audio or video file provided"}), 400

            # Once we have the transcript, we pass it to the translation system
            target_lang = "English"  # You can modify this to take from request
            translated_text = translate_text(transcript, target_lang)

            return jsonify({
                "status": "success",
                "input_type": "video" if "video" in request.files else "audio",
                "transcript": transcript,
                "translated_text": translated_text
            }), 200

        return jsonify({"error": "Invalid file or no file provided"}), 400

    except Exception as e:
        app.logger.error(f'Audio/Video error: {str(e)}', exc_info=True)
        return jsonify({"error": str(e), "status": "error"}), 500


@app.route('/api/chatbot/ask', methods=['POST'])
@login_required
@limiter.limit("30 per minute")
def chatbot_ask():
    try:
        data = request.get_json()
        question = data.get('question', '').strip()
        
        if not question:
            return jsonify({'error': 'Please provide a question'}), 400

        # استخدم النموذج للحصول على الرد من Hugging Face
        response = generate_huggingface_response(question)

        return jsonify({
            'response': response
        })
        
    except Exception as e:
        app.logger.error(f'Chatbot ask error: {str(e)}', exc_info=True)
        return jsonify({'error': str(e)}), 500
    
@app.route('/upload_video', methods=['POST'])
@login_required
@limiter.limit("5 per minute")
def upload_video():
    client_ip = get_remote_address()
    video_filepath = None
    audio_filepath = None
    try:
        if 'video' not in request.files:
            app.logger.error("No video file provided")
            return jsonify({'error': 'No video file provided'}), 400

        video_file = request.files['video']
        target_lang = request.form.get('target_lang', 'English')

        if video_file.filename == '':
            app.logger.error("Empty video file")
            return jsonify({'error': 'Empty video file'}), 400

        # حفظ الفيديو مؤقتًا
        filename = secure_filename(video_file.filename)
        video_filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        video_file.save(video_filepath)

        app.logger.info(f"Video file saved to: {video_filepath}")

        # معالجة الفيديو
        audio_filepath = os.path.splitext(video_filepath)[0] + ".wav"
        ffmpeg_cmd = _local_ffmpeg if os.path.exists(_local_ffmpeg) and os.access(_local_ffmpeg, os.X_OK) else "ffmpeg"
        ffmpeg.input(video_filepath).output(audio_filepath).run(cmd=ffmpeg_cmd)

        # تحويل الصوت إلى نص باستخدام whisper
        model = get_whisper_model()
        result = model.transcribe(audio_filepath)
        text = result['text']

        if not text:
            app.logger.error("No speech detected in the video")
            return jsonify({'error': 'No speech detected'}), 400

        app.logger.info(f"Extracted text from video: {text}")

        # ترجمته
        detected_lang = detect_language(text)
        translated_text = translate_text(text, target_lang)

        # حفظ الترجمة في قاعدة البيانات
        translation = Translation(
            user_id=current_user.id,
            original_text=text,
            detected_language=detected_lang,
            target_language=target_lang,
            translated_text=translated_text,
            is_favorite=False
        )
        db.session.add(translation)
        db.session.commit()

        return jsonify({
            'original_text': text,
            'detected_language': detected_lang,
            'target_language': target_lang,
            'translated_text': translated_text,
            'history_id': translation.id
        })

    except Exception as e:
        app.logger.error(f"Error processing video: {str(e)}", exc_info=True)
        return jsonify({'error': 'Video processing failed'}), 500

    finally:
        # تنظيف الملفات المؤقتة
        if video_filepath and os.path.exists(video_filepath):
            try:
                os.remove(video_filepath)
            except Exception as e:
                app.logger.warning(f"Failed to delete video file: {str(e)}")

        if audio_filepath and os.path.exists(audio_filepath):
            try:
                os.remove(audio_filepath)
            except Exception as e:
                app.logger.warning(f"Failed to delete audio file: {str(e)}")


@app.route('/statistics', methods=['POST'])
@login_required
@limiter.limit("20 per minute")  # Allow 20 statistics requests per minute per IP
def get_statistics():
    """Get statistics about the text"""
    try:
        data = request.get_json()
        text = data.get('text', '')
        
        if not text:
            return jsonify({'error': 'No text provided'}), 400
        
        words = text.split()
        characters = len(text)
        characters_no_spaces = len(text.replace(' ', ''))
        sentences = text.count('.') + text.count('!') + text.count('?')
        
        return jsonify({
            'word_count': len(words),
            'character_count': characters,
            'character_count_no_spaces': characters_no_spaces,
            'sentence_count': sentences,
            'average_word_length': sum(len(word) for word in words) / len(words) if words else 0
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/dashboard', methods=['GET'])
@login_required
def dashboard():
    """User dashboard page"""
    return render_template('dashboard.html', user=current_user)

@app.route('/api/dashboard/stats', methods=['GET'])
@login_required
@limiter.limit("30 per minute")  # Allow 30 dashboard requests per minute per IP
def get_dashboard_stats():
    """Get dashboard statistics for current user"""
    try:
        from sqlalchemy import func, extract
        from datetime import datetime, timedelta
        
        user_id = current_user.id
        
        # Total translations
        total_translations = Translation.query.filter_by(user_id=user_id).count()
        
        # Translations this month
        now = datetime.utcnow()
        first_day_of_month = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        this_month = Translation.query.filter(
            Translation.user_id == user_id,
            Translation.timestamp >= first_day_of_month
        ).count()
        
        # Translations this week
        week_ago = now - timedelta(days=7)
        this_week = Translation.query.filter(
            Translation.user_id == user_id,
            Translation.timestamp >= week_ago
        ).count()
        
        # Translations today
        today_start = now.replace(hour=0, minute=0, second=0, microsecond=0)
        today = Translation.query.filter(
            Translation.user_id == user_id,
            Translation.timestamp >= today_start
        ).count()
        
        # Favorite translations count
        favorites_count = Translation.query.filter_by(
            user_id=user_id,
            is_favorite=True
        ).count()
        
        # Language distribution (target languages)
        language_stats = db.session.query(
            Translation.target_language,
            func.count(Translation.id).label('count')
        ).filter_by(user_id=user_id).group_by(
            Translation.target_language
        ).order_by(func.count(Translation.id).desc()).limit(10).all()
        
        language_distribution = {lang: count for lang, count in language_stats if lang}
        
        # Activity over last 7 days
        activity_data = []
        for i in range(6, -1, -1):  # Last 7 days including today
            day_start = (now - timedelta(days=i)).replace(hour=0, minute=0, second=0, microsecond=0)
            day_end = day_start + timedelta(days=1)
            day_count = Translation.query.filter(
                Translation.user_id == user_id,
                Translation.timestamp >= day_start,
                Translation.timestamp < day_end
            ).count()
            activity_data.append({
                'date': day_start.strftime('%Y-%m-%d'),
                'day': day_start.strftime('%a'),
                'count': day_count
            })
        
        # Most common language pairs
        language_pairs = db.session.query(
            Translation.detected_language,
            Translation.target_language,
            func.count(Translation.id).label('count')
        ).filter_by(user_id=user_id).group_by(
            Translation.detected_language,
            Translation.target_language
        ).order_by(func.count(Translation.id).desc()).limit(5).all()
        
        top_language_pairs = [
            {'from': pair[0] or 'Unknown', 'to': pair[1], 'count': pair[2]}
            for pair in language_pairs
        ]
        
        # Average text length
        avg_length = db.session.query(
            func.avg(func.length(Translation.original_text))
        ).filter_by(user_id=user_id).scalar() or 0
        
        # Total characters translated
        total_chars = db.session.query(
            func.sum(func.length(Translation.original_text))
        ).filter_by(user_id=user_id).scalar() or 0
        
        return jsonify({
            'total_translations': total_translations,
            'this_month': this_month,
            'this_week': this_week,
            'today': today,
            'favorites_count': favorites_count,
            'language_distribution': language_distribution,
            'activity_timeline': activity_data,
            'top_language_pairs': top_language_pairs,
            'average_text_length': round(float(avg_length), 1),
            'total_characters': int(total_chars),
            'account_created': current_user.created_at.isoformat() if current_user.created_at else None,
            'last_login': current_user.last_login.isoformat() if current_user.last_login else None
        })
    except Exception as e:
        app.logger.error(f'Error getting dashboard stats for user {current_user.email}: {str(e)}', exc_info=True)
        return jsonify({'error': str(e)}), 500

@app.route('/history', methods=['GET'])
@login_required
@limiter.limit("30 per minute")  # Allow 30 history requests per minute per IP
def get_history():
    """Get translation history for current user"""
    try:
        limit = request.args.get('limit', 50, type=int)
        # Return history in reverse order (newest first) for current user only
        translations = Translation.query.filter_by(user_id=current_user.id).order_by(Translation.timestamp.desc()).limit(limit).all()
        history_list = [t.to_dict() for t in translations]
        total = Translation.query.filter_by(user_id=current_user.id).count()
        return jsonify({
            'history': history_list,
            'total': total
        })
    except Exception as e:
        app.logger.error(f'Error getting history for user {current_user.email}: {str(e)}', exc_info=True)
        return jsonify({'error': str(e)}), 500

@app.route('/history/<int:history_id>', methods=['DELETE'])
@login_required
@limiter.limit("10 per minute")  # Allow 10 delete operations per minute per IP
def delete_history_item(history_id):
    """Delete a specific history item (only if owned by current user)"""
    try:
        translation = Translation.query.filter_by(id=history_id, user_id=current_user.id).first_or_404()
        db.session.delete(translation)
        db.session.commit()
        app.logger.info(f'Translation {history_id} deleted by user {current_user.email}')
        return jsonify({'message': 'Deleted', 'id': history_id})
    except Exception as e:
        db.session.rollback()
        app.logger.warning(f'Failed to delete translation {history_id} for user {current_user.email}: {str(e)}')
        return jsonify({'error': 'History item not found'}), 404

@app.route('/history/clear', methods=['DELETE'])
@login_required
@limiter.limit("5 per minute")  # Allow 5 clear operations per minute per IP (destructive action)
def clear_history():
    """Clear all translation history for current user"""
    try:
        count = Translation.query.filter_by(user_id=current_user.id).count()
        Translation.query.filter_by(user_id=current_user.id).delete()
        db.session.commit()
        app.logger.info(f'History cleared for user {current_user.email}: {count} translations deleted')
        return jsonify({'message': 'History cleared', 'deleted_count': count})
    except Exception as e:
        db.session.rollback()
        app.logger.error(f'Error clearing history for user {current_user.email}: {str(e)}', exc_info=True)
        return jsonify({'error': str(e)}), 500

@app.route('/history/<int:history_id>/favorite', methods=['POST'])
@login_required
@limiter.limit("20 per minute")  # Allow 20 favorite toggles per minute per IP
def toggle_favorite(history_id):
    """Toggle favorite status of a translation (only if owned by current user)"""
    try:
        translation = Translation.query.filter_by(id=history_id, user_id=current_user.id).first_or_404()
        translation.is_favorite = not translation.is_favorite
        db.session.commit()
        app.logger.debug(f'Favorite toggled for translation {history_id} by user {current_user.email}')
        return jsonify({
            'message': 'Favorite toggled',
            'id': history_id,
            'is_favorite': translation.is_favorite
        })
    except Exception as e:
        db.session.rollback()
        app.logger.warning(f'Failed to toggle favorite for translation {history_id} by user {current_user.email}: {str(e)}')
        return jsonify({'error': 'Translation not found'}), 404

@app.route('/translate_batch', methods=['POST'])
@login_required
@limiter.limit("3 per minute")  # Allow 3 batch translations per minute per IP
def translate_batch():
    """Translate multiple texts at once"""
    try:
        data = request.get_json()
        texts = data.get('texts', [])
        target_lang = data.get('target_lang', 'English')
        
        if not texts or not isinstance(texts, list):
            return jsonify({'error': 'Please provide a list of texts'}), 400
        
        if len(texts) > 50:  # Limit batch size
            return jsonify({'error': 'Maximum 50 texts per batch'}), 400
        
        results = []
        for text in texts:
            if not text.strip():
                continue
            try:
                detected_lang = detect_language(text)
                translated_text = translate_text(text, target_lang)
                results.append({
                    'original': text,
                    'detected_language': detected_lang,
                    'translated': translated_text
                })
            except Exception as e:
                results.append({
                    'original': text,
                    'detected_language': 'Unknown',
                    'translated': f'Error: {str(e)}'
                })
        
        # Save batch to database with user_id
        for result in results:
            translation = Translation(
                user_id=current_user.id,
                original_text=result['original'],
                detected_language=result['detected_language'],
                target_language=target_lang,
                translated_text=result['translated'],
                is_favorite=False
            )
            db.session.add(translation)
        
        db.session.commit()
        app.logger.info(f'Batch translation saved for user {current_user.email}: {len(results)} translations')
        
        # Keep only last 100 translations per user (optional cleanup)
        user_translations_count = Translation.query.filter_by(user_id=current_user.id).count()
        if user_translations_count > 100:
            oldest = Translation.query.filter_by(user_id=current_user.id).order_by(Translation.timestamp.asc()).limit(user_translations_count - 100).all()
            for old_translation in oldest:
                db.session.delete(old_translation)
            db.session.commit()
            app.logger.info(f'Cleaned up old batch translations for user {current_user.email}')
        
        return jsonify({
            'target_language': target_lang,
            'count': len(results),
            'results': results
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/export', methods=['POST'])
@login_required
@limiter.limit("10 per minute")  # Allow 10 exports per minute per IP
def export_translation():
    """Export translation in different formats"""
    try:
        data = request.get_json()
        original = data.get('original_text', '')
        translated = data.get('translated_text', '')
        source_lang = data.get('source_language', 'Unknown')
        target_lang = data.get('target_language', 'Unknown')
        format_type = data.get('format', 'txt')  # txt, json, csv
        
        if not original or not translated:
            return jsonify({'error': 'No translation data to export'}), 400
        
        if format_type == 'txt':
            content = f"Original Text ({source_lang}):\n{original}\n\nTranslated Text ({target_lang}):\n{translated}\n\nExported: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            mimetype = 'text/plain'
            filename = 'translation.txt'
        elif format_type == 'json':
            content = json.dumps({
                'original_text': original,
                'translated_text': translated,
                'source_language': source_lang,
                'target_language': target_lang,
                'exported_at': datetime.now().isoformat()
            }, indent=2, ensure_ascii=False)
            mimetype = 'application/json'
            filename = 'translation.json'
        elif format_type == 'csv':
            output = io.StringIO()
            writer = csv.writer(output)
            writer.writerow(['Field', 'Value'])
            writer.writerow(['Source Language', source_lang])
            writer.writerow(['Target Language', target_lang])
            writer.writerow(['Original Text', original])
            writer.writerow(['Translated Text', translated])
            writer.writerow(['Exported At', datetime.now().isoformat()])
            content = output.getvalue()
            mimetype = 'text/csv'
            filename = 'translation.csv'
        elif format_type == 'pdf':
            # Generate PDF using reportlab
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.lib.enums import TA_LEFT, TA_CENTER
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            
            try:
                # Create PDF in memory
                buffer = io.BytesIO()
                doc = SimpleDocTemplate(buffer, pagesize=A4, 
                                      rightMargin=72, leftMargin=72,
                                      topMargin=72, bottomMargin=18)
                
                # Container for the 'Flowable' objects
                story = []
                
                # Define styles
                styles = getSampleStyleSheet()
                title_style = ParagraphStyle(
                    'CustomTitle',
                    parent=styles['Heading1'],
                    fontSize=18,
                    textColor='#2c3e50',
                    spaceAfter=30,
                    alignment=TA_CENTER
                )
                heading_style = ParagraphStyle(
                    'CustomHeading',
                    parent=styles['Heading2'],
                    fontSize=14,
                    textColor='#34495e',
                    spaceAfter=12,
                    spaceBefore=20
                )
                normal_style = ParagraphStyle(
                    'CustomNormal',
                    parent=styles['Normal'],
                    fontSize=11,
                    leading=16,
                    spaceAfter=12
                )
                info_style = ParagraphStyle(
                    'CustomInfo',
                    parent=styles['Normal'],
                    fontSize=9,
                    textColor='#7f8c8d',
                    spaceAfter=6
                )
                
                # Title
                story.append(Paragraph("Translation Document", title_style))
                story.append(Spacer(1, 0.2*inch))
                
                # Export info
                export_info = f"Exported on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
                story.append(Paragraph(export_info, info_style))
                story.append(Spacer(1, 0.3*inch))
                
                # Original Text Section
                story.append(Paragraph(f"Original Text ({source_lang})", heading_style))
                story.append(Paragraph(original.replace('\n', '<br/>'), normal_style))
                story.append(Spacer(1, 0.3*inch))
                
                # Translated Text Section
                story.append(Paragraph(f"Translated Text ({target_lang})", heading_style))
                story.append(Paragraph(translated.replace('\n', '<br/>'), normal_style))
                
                # Build PDF
                doc.build(story)
                
                # Get PDF content
                buffer.seek(0)
                content = buffer.read()
                buffer.close()
                
                mimetype = 'application/pdf'
                filename = 'translation.pdf'
                
            except Exception as pdf_error:
                app.logger.error(f'PDF generation error: {str(pdf_error)}', exc_info=True)
                return jsonify({'error': f'Failed to generate PDF: {str(pdf_error)}'}), 500
                
        elif format_type == 'docx':
            # Generate DOCX using python-docx
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            try:
                # Create document
                doc = Document()
                
                # Set document margins
                sections = doc.sections
                for section in sections:
                    section.top_margin = Inches(0.75)
                    section.bottom_margin = Inches(0.75)
                    section.left_margin = Inches(0.75)
                    section.right_margin = Inches(0.75)
                
                # Title
                title = doc.add_heading('Translation Document', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Export info
                export_para = doc.add_paragraph()
                export_run = export_para.add_run(f"Exported on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                export_run.font.size = Pt(9)
                export_run.font.color.rgb = RGBColor(127, 140, 141)
                export_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add spacing
                doc.add_paragraph()
                
                # Original Text Section
                doc.add_heading(f'Original Text ({source_lang})', level=2)
                original_para = doc.add_paragraph()
                original_run = original_para.add_run(original)
                original_run.font.size = Pt(11)
                
                # Add spacing
                doc.add_paragraph()
                
                # Translated Text Section
                doc.add_heading(f'Translated Text ({target_lang})', level=2)
                translated_para = doc.add_paragraph()
                translated_run = translated_para.add_run(translated)
                translated_run.font.size = Pt(11)
                
                # Save to BytesIO
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                content = buffer.read()
                buffer.close()
                
                mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                filename = 'translation.docx'
                
            except Exception as docx_error:
                app.logger.error(f'DOCX generation error: {str(docx_error)}', exc_info=True)
                return jsonify({'error': f'Failed to generate DOCX: {str(docx_error)}'}), 500
                
        else:
            return jsonify({'error': 'Invalid format. Use: txt, json, csv, pdf, or docx'}), 400
        
        return Response(
            content,
            mimetype=mimetype,
            headers={'Content-Disposition': f'attachment; filename={filename}'}
        )
    except Exception as e:
        app.logger.error(f'Export error: {str(e)}', exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/text_to_speech', methods=['POST'])
@login_required
@limiter.limit("10 per minute")
def text_to_speech():
    try:
        data = request.get_json() or {}
        text = (data.get('text') or '').strip()
        lang = (data.get('lang') or '').strip()

        if not text:
            return jsonify({'error': 'No text provided'}), 400

        # Accept either language name (English) or code (en)
        if not lang:
            lang_code = 'en'
        elif lang in LANG_CODES:
            lang_code = LANG_CODES[lang]
        elif lang in LANGUAGES:
            lang_code = lang
        else:
            name_to_code = {v: k for k, v in LANGUAGES.items()}
            lang_code = name_to_code.get(lang, 'en')

        tts = gTTS(text=text, lang=lang_code, slow=False)
        audio_buffer = io.BytesIO()
        tts.write_to_fp(audio_buffer)
        audio_buffer.seek(0)

        return send_file(
            audio_buffer,
            mimetype='audio/mpeg',
            as_attachment=False,
            download_name='tts.mp3'
        )
    except Exception as e:
        app.logger.error(f'TTS error: {str(e)}', exc_info=True)
        return jsonify({'error': 'Text-to-speech failed'}), 500

# ========== ADMIN DASHBOARD ROUTES ==========

@app.route('/admin', methods=['GET'])
@admin_required
def admin_dashboard():
    """Admin dashboard page"""
    # Ensure we have the latest user object with is_admin
    try:
        db.session.refresh(current_user)
    except Exception:
        pass
    return render_template('admin_dashboard.html', user=current_user)

@app.route('/api/admin/stats', methods=['GET'])
@admin_required
@limiter.limit("30 per minute")
def get_admin_stats():
    """Get admin dashboard statistics"""
    try:
        from sqlalchemy import func
        from datetime import datetime, timedelta
        
        # Total users
        total_users = User.query.count()
        
        # Total translations
        total_translations = Translation.query.count()
        
        # Active users (logged in last 30 days)
        thirty_days_ago = datetime.utcnow() - timedelta(days=30)
        active_users = User.query.filter(
            User.last_login >= thirty_days_ago
        ).count()
        
        # New users this month
        now = datetime.utcnow()
        first_day_of_month = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        new_users_month = User.query.filter(
            User.created_at >= first_day_of_month
        ).count()
        
        # New users this week
        week_ago = now - timedelta(days=7)
        new_users_week = User.query.filter(
            User.created_at >= week_ago
        ).count()
        
        # Translations this month
        translations_month = Translation.query.filter(
            Translation.timestamp >= first_day_of_month
        ).count()
        
        # Translations this week
        translations_week = Translation.query.filter(
            Translation.timestamp >= week_ago
        ).count()
        
        # Verified users
        verified_users = User.query.filter_by(email_verified=True).count()
        
        # Admin users
        admin_users = User.query.filter_by(is_admin=True).count()
        
        # Language distribution (all users)
        language_stats = db.session.query(
            Translation.target_language,
            func.count(Translation.id).label('count')
        ).group_by(Translation.target_language).order_by(
            func.count(Translation.id).desc()
        ).limit(10).all()
        
        language_distribution = {lang: count for lang, count in language_stats if lang}
        
        # Users by provider
        provider_stats = db.session.query(
            User.provider,
            func.count(User.id).label('count')
        ).group_by(User.provider).all()
        
        provider_distribution = {provider: count for provider, count in provider_stats}
        
        # Activity over last 7 days
        activity_data = []
        for i in range(6, -1, -1):
            day_start = (now - timedelta(days=i)).replace(hour=0, minute=0, second=0, microsecond=0)
            day_end = day_start + timedelta(days=1)
            day_count = Translation.query.filter(
                Translation.timestamp >= day_start,
                Translation.timestamp < day_end
            ).count()
            activity_data.append({
                'date': day_start.strftime('%Y-%m-%d'),
                'day': day_start.strftime('%a'),
                'count': day_count
            })
        
        return jsonify({
            'total_users': total_users,
            'total_translations': total_translations,
            'active_users': active_users,
            'new_users_month': new_users_month,
            'new_users_week': new_users_week,
            'translations_month': translations_month,
            'translations_week': translations_week,
            'verified_users': verified_users,
            'admin_users': admin_users,
            'language_distribution': language_distribution,
            'provider_distribution': provider_distribution,
            'activity_timeline': activity_data
        })
    except Exception as e:
        app.logger.error(f'Admin stats error: {str(e)}', exc_info=True)
        return jsonify({'error': str(e)}), 500

@app.route('/api/admin/users', methods=['GET'])
@admin_required
@limiter.limit("30 per minute")
def get_admin_users():
    """Get list of all users for admin dashboard"""
    try:
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 20, type=int)
        search = request.args.get('search', '', type=str)
        sort_by = request.args.get('sort_by', 'created_at', type=str)
        sort_order = request.args.get('sort_order', 'desc', type=str)
        
        query = User.query
        
        # Search filter
        if search:
            query = query.filter(
                db.or_(
                    User.email.ilike(f'%{search}%'),
                    User.name.ilike(f'%{search}%'),
                    User.username.ilike(f'%{search}%')
                )
            )
        
        # Sorting
        if hasattr(User, sort_by):
            order_column = getattr(User, sort_by)
            if sort_order == 'desc':
                query = query.order_by(order_column.desc())
            else:
                query = query.order_by(order_column.asc())
        
        # Pagination
        pagination = query.paginate(page=page, per_page=per_page, error_out=False)
        users = pagination.items
        
        users_data = []
        for user in users:
            translations_count = Translation.query.filter_by(user_id=user.id).count()
            users_data.append({
                'id': user.id,
                'email': user.email,
                'username': user.username,
                'name': user.name,
                'provider': user.provider,
                'email_verified': user.email_verified,
                'is_admin': user.is_admin,
                'created_at': user.created_at.isoformat() if user.created_at else None,
                'last_login': user.last_login.isoformat() if user.last_login else None,
                'translations_count': translations_count,
                'picture': user.picture
            })
        
        return jsonify({
            'users': users_data,
            'total': pagination.total,
            'pages': pagination.pages,
            'current_page': page,
            'per_page': per_page
        })
    except Exception as e:
        app.logger.error(f'Admin users error: {str(e)}', exc_info=True)
        return jsonify({'error': str(e)}), 500

@app.route('/api/admin/translations', methods=['GET'])
@admin_required
@limiter.limit("30 per minute")
def get_admin_translations():
    """Get list of all translations for admin dashboard"""
    try:
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 20, type=int)
        search = request.args.get('search', '', type=str)
        user_id = request.args.get('user_id', type=int)
        
        query = Translation.query
        
        # Filters
        if search:
            query = query.filter(
                db.or_(
                    Translation.original_text.ilike(f'%{search}%'),
                    Translation.translated_text.ilike(f'%{search}%')
                )
            )
        
        if user_id:
            query = query.filter_by(user_id=user_id)
        
        # Order by timestamp desc
        query = query.order_by(Translation.timestamp.desc())
        
        # Pagination
        pagination = query.paginate(page=page, per_page=per_page, error_out=False)
        translations = pagination.items
        
        translations_data = []
        for trans in translations:
            user = User.query.get(trans.user_id)
            translations_data.append({
                'id': trans.id,
                'user_id': trans.user_id,
                'user_email': user.email if user else 'Unknown',
                'original_text': trans.original_text[:100] + '...' if len(trans.original_text) > 100 else trans.original_text,
                'translated_text': trans.translated_text[:100] + '...' if len(trans.translated_text) > 100 else trans.translated_text,
                'detected_language': trans.detected_language,
                'target_language': trans.target_language,
                'timestamp': trans.timestamp.isoformat(),
                'is_favorite': trans.is_favorite
            })
        
        return jsonify({
            'translations': translations_data,
            'total': pagination.total,
            'pages': pagination.pages,
            'current_page': page,
            'per_page': per_page
        })
    except Exception as e:
        app.logger.error(f'Admin translations error: {str(e)}', exc_info=True)
        return jsonify({'error': str(e)}), 500

# Chatbot endpoints
@app.route('/api/chatbot/context', methods=['POST'])
@login_required
@limiter.limit("20 per minute")
def get_chatbot_context():
    """Detect topic and generate context for chatbot"""
    try:
        data = request.get_json()
        text = data.get('text', '').strip().lower()
        
        if not text or len(text) < 10:
            return jsonify({'error': 'Text too short for analysis'}), 400
        
        # Topic detection based on keywords
        topics_keywords = {
            'Economics': ['economy', 'economic', 'market', 'trade', 'gdp', 'inflation', 'currency', 'financial', 'business', 'stock', 'investment', 'revenue', 'profit', 'cost', 'price', 'supply', 'demand'],
            'Science': ['scientific', 'research', 'experiment', 'hypothesis', 'theory', 'study', 'discovery', 'laboratory', 'chemical', 'physics', 'biology', 'chemistry'],
            'Technology': ['technology', 'software', 'computer', 'digital', 'internet', 'app', 'programming', 'code', 'algorithm', 'data', 'artificial intelligence', 'ai', 'machine learning'],
            'Health': ['health', 'medical', 'medicine', 'doctor', 'patient', 'treatment', 'disease', 'symptom', 'diagnosis', 'hospital', 'therapy'],
            'Education': ['education', 'school', 'student', 'teacher', 'learning', 'university', 'study', 'course', 'curriculum', 'academic'],
            'Politics': ['political', 'government', 'election', 'policy', 'democracy', 'president', 'parliament', 'vote', 'law', 'legislation'],
            'History': ['history', 'historical', 'ancient', 'past', 'event', 'war', 'empire', 'civilization', 'culture'],
            'Sports': ['sport', 'sports', 'game', 'games', 'team', 'teams', 'player', 'players', 'match', 'matches', 'championship', 'tournament', 'tournaments', 'athlete', 'athletes', 'competition', 'competitions', 'football', 'soccer', 'basketball', 'baseball', 'tennis', 'volleyball', 'hockey', 'cricket', 'rugby', 'golf', 'swimming', 'running', 'racing', 'cycling', 'boxing', 'wrestling', 'score', 'scoring', 'goal', 'goals', 'field', 'pitch', 'court', 'stadium', 'referee', 'coach', 'training', 'practice', 'league', 'cup', 'trophy', 'victory', 'defeat', 'win', 'lose', 'offense', 'defense', 'attack', 'defend'],
            'Art': ['art', 'artist', 'painting', 'music', 'literature', 'creative', 'design', 'culture', 'museum'],
            'Environment': ['environment', 'climate', 'pollution', 'nature', 'green', 'sustainability', 'renewable', 'energy', 'carbon']
        }
        
        detected_topic = 'General'
        max_matches = 0
        
        for topic, keywords in topics_keywords.items():
            matches = sum(1 for keyword in keywords if keyword in text)
            if matches > max_matches:
                max_matches = matches
                detected_topic = topic
        
        # Generate context based on topic
        context_info = generate_topic_context(detected_topic, text)
        
        app.logger.info(f'Chatbot context requested by {current_user.email}, topic: {detected_topic}')
        
        return jsonify({
            'topic': detected_topic,
            'context': context_info
        })
        
    except Exception as e:
        app.logger.error(f'Chatbot context error: {str(e)}', exc_info=True)
        return jsonify({'error': str(e)}), 500

def generate_topic_context(topic, text):
    """Generate helpful context information based on topic"""
    context_templates = {
        'Economics': """Economics deals with the production, distribution, and consumption of goods and services. Key concepts include supply and demand, market forces, monetary policy, and trade. Understanding economic terms in context helps ensure accurate translation of financial and business content.""",
        'Science': """Scientific translations require precision with technical terms, units of measurement, and research terminology. Pay attention to scientific names, formulas, and standard scientific notation. Accuracy is crucial for maintaining scientific validity.""",
        'Technology': """Technology translations often involve technical jargon, product names, and industry-specific terminology. Software terms, programming concepts, and technical specifications need careful translation to preserve meaning across languages.""",
        'Health': """Medical and health translations require accuracy for safety and clarity. Medical terms, symptoms, treatments, and procedures should be translated using proper medical terminology. Context is critical in healthcare communication.""",
        'Education': """Educational content translation involves academic terminology, curriculum terms, and pedagogical concepts. Consider the educational level and cultural context when translating educational materials.""",
        'Politics': """Political translations require understanding of governmental structures, policies, and political systems. Political terminology can vary significantly between cultures and legal systems.""",
        'History': """Historical translations involve names, dates, events, and cultural contexts. Historical accuracy and proper handling of names and places are important for maintaining historical integrity.""",
        'Sports': """Sports translations include game terminology, rules, player names, and event descriptions. Popular sports like football (soccer), basketball, tennis, and others have specific terminology. Sports terms often have established translations, but cultural contexts matter - for example, 'football' refers to soccer in most countries but American football in the US. When translating sports content, consider the target audience's preferred terminology and maintain consistency with established sports vocabulary.""",
        'Art': """Art and cultural translations require sensitivity to creative expression and cultural nuances. Artistic terminology, styles, and cultural references need careful handling to preserve artistic intent.""",
        'Environment': """Environmental translations involve climate science, sustainability terms, and ecological concepts. Accurate translation of environmental data and policies is important for global environmental communication."""
    }
    
    base_context = context_templates.get(topic, 
        """This is general content. For accurate translation, consider the context, audience, and cultural nuances. Technical terms should be verified for domain-specific meanings.""")
    
    # Add specific tips based on text length
    if len(text) > 200:
        base_context += "\n\nTip: This appears to be a longer text. Consider reviewing key terms and maintaining consistency throughout the translation."
    
    return base_context

def generate_ai_chatbot_response(question, topic, context, api_key, conversation_history=None):
    """Generate intelligent AI-powered response using OpenAI with conversation history"""
    try:
        # Initialize OpenAI client (v1.40.0 compatible with httpx 0.27.2)
        client = OpenAI(api_key=api_key)
        
        # Build a comprehensive prompt that includes context about the translation topic
        system_prompt = f"""You are a helpful translation assistant chatbot. The user is working on translating content related to {topic}.

Context about the topic:
{context}

Your role is to:
1. Provide helpful, accurate, and contextually relevant answers about the topic
2. Help with translation-related questions and terminology
3. Offer insights that can assist with understanding the content being translated
4. Be concise but informative
5. If asked about translation, provide practical translation tips specific to {topic}
6. Remember previous questions and answers in the conversation to provide context-aware follow-up responses

Respond naturally and conversationally. If you don't know something specific, say so honestly."""

        # Build messages array with system prompt and conversation history
        messages = [{"role": "system", "content": system_prompt}]
        
        # Add conversation history if available
        if conversation_history:
            app.logger.debug(f'Processing conversation history with {len(conversation_history)} messages')
            # Filter out system messages and add user/assistant messages
            for idx, msg in enumerate(conversation_history):
                if isinstance(msg, dict) and 'role' in msg and 'content' in msg:
                    role = msg.get('role', '').lower()
                    content = msg.get('content', '').strip()
                    if role in ['user', 'assistant'] and content:
                        # Clean up content (remove markdown formatting from welcome message if needed)
                        # but keep the content meaningful
                        messages.append({"role": role, "content": content})
                        app.logger.debug(f'  Added message {idx}: {role} - {content[:50]}...')
        else:
            app.logger.debug('No conversation history provided')
        
        # Add the current question
        messages.append({"role": "user", "content": question})
        
        app.logger.debug(f'Sending {len(messages)} messages to OpenAI (including system prompt)')
        
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=messages,
            temperature=0.7,
            max_tokens=500
        )
        
        result = response.choices[0].message.content.strip()
        app.logger.debug(f'Received response from OpenAI: {len(result)} characters')
        
        return result
        
    except Exception as e:
        app.logger.error(f'OpenAI API error: {str(e)}')
        raise


def generate_chatbot_response(question, topic, context):
    """Generate helpful response based on question and topic"""
    question_lower = question.lower()
    
    # Topic-specific knowledge base and response generation
    if topic == 'Sports':
        # Football/Soccer specific responses
        if any(word in question_lower for word in ['football', 'soccer', 'ball']):
            if any(word in question_lower for word in ['rules', 'rule', 'how to play', 'played']):
                return """Football (soccer) is played with two teams of 11 players each. The objective is to score goals by getting the ball into the opponent's net. Players cannot use their hands or arms (except the goalkeeper in their penalty area). The game is played in two halves of 45 minutes each. Offside, fouls, and cards are important rules to understand when translating football content."""
            elif any(word in question_lower for word in ['positions', 'position', 'player', 'role']):
                return """Common football positions include: goalkeeper (GK), defenders (center-back, full-back, wing-back), midfielders (defensive, central, attacking), and forwards (striker, winger). Each position has specific responsibilities and terminology that may vary slightly in translation."""
            elif any(word in question_lower for word in ['terms', 'terminology', 'words', 'vocabulary']):
                return """Key football terms include: goal, penalty, corner kick, free kick, offside, foul, yellow card, red card, substitution, injury time, extra time. When translating, consider regional variations - for example, 'pitch' (UK) vs 'field' (US), 'match' (UK) vs 'game' (US)."""
            elif any(word in question_lower for word in ['translate', 'translation', 'how to']):
                return """When translating football content:\n\n• Use established sports terminology in the target language\n• Preserve proper nouns (player names, team names, stadium names)\n• Maintain consistency with official translations (FIFA, UEFA)\n• Consider cultural context (American vs British English)\n• Keep technical terms accurate (offside, penalty, etc.)\n• Verify team and league names have standard translations"""
            else:
                return """Football (soccer) is the world's most popular sport, played by billions of people. Key aspects include team strategy, player positions, rules, and terminology. When translating football content, consider the audience's familiarity with the sport and use established translations for technical terms.\n\nWould you like to know more about football rules, positions, terminology, or translation tips?"""
        
        # General sports questions
        elif any(word in question_lower for word in ['sport', 'sports', 'athletic']):
            if any(word in question_lower for word in ['popular', 'common', 'types']):
                return """Popular sports include: football (soccer), basketball, tennis, volleyball, cricket, rugby, baseball, American football, hockey, and many others. Each sport has its own terminology, rules, and cultural significance that should be considered in translation."""
            elif any(word in question_lower for word in ['translate', 'translation']):
                return """Sports translation tips:\n\n• Maintain consistency with established sports terminology\n• Preserve proper nouns (player names, team names, venues)\n• Consider regional variations (e.g., football vs soccer)\n• Use official translations when available (FIFA, IOC, etc.)\n• Understand the rules and context of the sport\n• Keep technical terms accurate (foul, penalty, etc.)"""
            else:
                return """Sports encompass a wide range of physical activities and competitions. Each sport has unique terminology, rules, and cultural significance. When translating sports content, accuracy with technical terms and understanding of the sport's context are essential for maintaining meaning."""
        
        # General response for sports topic
        else:
            return """In sports translation, accuracy with terminology is crucial. Consider:\n\n• Each sport has specific technical terms that should use established translations\n• Player names, team names, and venue names should be preserved or use standard translations\n• Rules and regulations need precise translation to maintain accuracy\n• Cultural context matters - the same sport may have different terminology in different regions\n\nWhat specific aspect of sports would you like to know more about?"""
    
    elif topic == 'Economics':
        if any(word in question_lower for word in ['market', 'trading', 'stock']):
            return """Market-related translations require precision with financial terminology. Key concepts include: market trends, trading volumes, stock prices, market capitalization, and economic indicators. Use established financial terminology and maintain consistency with regulatory standards."""
        elif any(word in question_lower for word in ['inflation', 'economy', 'gdp']):
            return """Economic indicators like GDP, inflation rates, and economic growth require accurate translation. These terms have standard definitions and should be translated using official economic terminology. Consider the target audience's familiarity with economic concepts."""
        else:
            return """Economics involves the study of production, distribution, and consumption. Key areas include microeconomics (individual markets) and macroeconomics (overall economy). When translating economic content, use established terminology and maintain accuracy with statistical data and financial figures."""
    
    elif topic == 'Technology':
        if any(word in question_lower for word in ['software', 'app', 'programming']):
            return """Software and programming translations require technical accuracy. Programming terms, API names, and technical specifications should be carefully translated or kept in English when appropriate. Consider the developer community's conventions in the target language."""
        elif any(word in question_lower for word in ['ai', 'artificial intelligence', 'machine learning']):
            return """AI and machine learning translations involve cutting-edge terminology that may not have established translations in all languages. Often, technical terms remain in English. Verify the target language's conventions for technical AI terminology."""
        else:
            return """Technology translations involve technical jargon, product specifications, and industry terminology. Maintain technical accuracy while ensuring the translation is understandable to the target audience."""
    
    # Default response pattern matching
    elif any(word in question_lower for word in ['what', 'explain', 'meaning']):
        return f"Based on {topic}:\n\n{context}\n\nIs there a specific aspect of {topic} you'd like me to explain in more detail?"
    
    elif any(word in question_lower for word in ['how', 'translate']):
        return f"For translating {topic} content:\n\n• Use domain-specific terminology\n• Maintain technical accuracy\n• Consider cultural context\n• Verify key terms\n• Keep consistency throughout\n\n{context}"
    
    elif any(word in question_lower for word in ['why', 'important']):
        return f"Understanding {topic} context is important for accurate translation because it ensures proper terminology, preserves technical meaning, and maintains professional standards. {context[:200]}..."
    
    else:
        return f"Regarding {topic}:\n\n{context}\n\nFeel free to ask me more specific questions about this topic. I can help with terminology, translation tips, or key concepts!"
@app.route('/api/ai/assistant', methods=['POST'])
def ai_assistant():
    try:
        data = request.get_json(silent=True) or {}

        text = (data.get("content") or "").strip()
        if not text:
            return jsonify({"error": "content is required"}), 400

        result = handle_request(text, meta=data)

        return jsonify({
            "status": "success",
            "input_type": "text",
            **result
        })

    except Exception as e:
        app.logger.error(str(e), exc_info=True)
        return jsonify({"error": str(e), "status": "error"}), 500

    """
    Unified AI Assistant Endpoint (Integrated System)

    Supports:
    1) JSON text:
       {
         "input_type": "text",
         "content": "...",
         "action": "summarize|rewrite|improve|translate|...",
         "target_lang": "Arabic|English|...",
         "meta": {...}
       }

    2) Multipart audio (form-data):
       - audio: file
       - action: optional
       - target_lang: optional
       - meta: optional (string/json later)
    """
    try:
        # -------------------------
        # Case A: multipart/form-data (audio)
        # -------------------------
        if request.content_type and "multipart/form-data" in request.content_type:
            if "audio" not in request.files:
                return jsonify({"status": "error", "error": "No audio file provided"}), 400

            audio_file = request.files["audio"]
            if not audio_file or audio_file.filename == "":
                return jsonify({"status": "error", "error": "Empty audio filename"}), 400

            # options / meta
            action = (request.form.get("action") or "").strip().lower() or None
            target_lang = request.form.get("target_lang", "English")
            # لاحقًا يمكن قراءة meta JSON من request.form.get("meta")
            meta = {"action": action, "target_lang": target_lang}

            filename = secure_filename(audio_file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            audio_file.save(filepath)

            try:
                transcript = transcribe_audio(filepath)
                if not transcript or transcript.startswith("Error") or transcript.startswith("Could not"):
                    return jsonify({"status": "error", "error": transcript or "Transcription failed"}), 400
            finally:
                if os.path.exists(filepath):
                    os.remove(filepath)

            core_result = handle_request(transcript, meta)

            return jsonify({
                "status": "success",
                "input_type": "voice",
                "transcript": transcript,
                "intent": core_result.get("intent"),
                "result": core_result
            }), 200

        # -------------------------
        # Case B: application/json (text)
        # -------------------------
        data = request.get_json(silent=True) or {}

        input_type = (data.get("input_type") or "text").strip().lower()
        content = (data.get("content") or "").strip()

        if input_type not in ["text", "voice"]:
            return jsonify({"status": "error", "error": "input_type must be text or voice"}), 400

        if not content:
            return jsonify({"status": "error", "error": "content is required"}), 400

        # Voice عبر JSON غير مدعوم هنا (منعًا للتعقيد)
        if input_type == "voice":
            return jsonify({
                "status": "error",
                "error": "For voice, send multipart/form-data with an 'audio' file."
            }), 400

        action = (data.get("action") or "").strip().lower() or None
        target_lang = data.get("target_lang", "English")

        # meta object اختياري من العميل
        user_meta = data.get("meta") or {}
        if not isinstance(user_meta, dict):
            user_meta = {}

        meta = {
            **user_meta,
            "action": action,
            "target_lang": target_lang
        }

        core_result = handle_request(content, meta)

        return jsonify({
            "status": "success",
            "input_type": "text",
            "intent": core_result.get("intent"),
            "result": core_result
        }), 200

    except Exception as e:
        app.logger.error(f'AI assistant error: {str(e)}', exc_info=True)
        return jsonify({"status": "error", "error": str(e)}), 500

if __name__ == '__main__':
    # Initialize database tables
    try:
        with app.app_context():
            db.create_all()
            app.logger.info("✅ Database tables created/verified")
            print("✅ Database tables created/verified")
    except Exception as e:
        app.logger.error(f'Database initialization error: {str(e)}', exc_info=True)
        print(f"❌ Database initialization error: {str(e)}")
        raise

    # Preload summarization model in background to avoid blocking startup
    try:
        if os.getenv("PRELOAD_SUMMARIZER", "true").lower() == "true":
            import threading

            def _preload():
                try:
                    get_summarizer()
                    app.logger.info("✅ Summarizer model preloaded")
                except Exception as e:
                    app.logger.warning(f"⚠️ Summarizer preload failed: {e}")

            threading.Thread(target=_preload, daemon=True).start()
        else:
            app.logger.info("ℹ️ Summarizer preload skipped (PRELOAD_SUMMARIZER=false)")
    except Exception as e:
        app.logger.warning(f"⚠️ Summarizer preload setup failed: {e}")
    
    # Production: use PORT from environment, default to 5000
    port = int(os.environ.get('PORT', 5000))
    # Only enable debug in development
    debug = os.environ.get('FLASK_ENV') == 'development'
    
    app.logger.info(f'Starting Flask server on port {port} (debug={debug})')
    app.run(host='0.0.0.0', port=port, debug=debug)


